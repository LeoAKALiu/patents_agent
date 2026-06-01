from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document

from backend.app.schemas import PatentClaim, PatentDocument, PatentSection, SectionType


HEADING_MAP: dict[str, SectionType] = {
    "摘要": SectionType.ABSTRACT,
    "说明书摘要": SectionType.ABSTRACT,
    "权利要求书": SectionType.CLAIMS,
    "权利要求": SectionType.CLAIMS,
    "说明书": SectionType.DESCRIPTION,
    "技术领域": SectionType.TECHNICAL_FIELD,
    "背景技术": SectionType.BACKGROUND,
    "发明内容": SectionType.SUMMARY,
    "实用新型内容": SectionType.SUMMARY,
    "附图说明": SectionType.DRAWINGS,
    "具体实施方式": SectionType.EMBODIMENTS,
    "具体实施例": SectionType.EMBODIMENTS,
}


def read_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf_text(path)
    if suffix == ".docx":
        return _read_docx_text(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return _read_plain_text(path)
    if suffix == ".xml":
        return _read_xml_text(path)
    raise ValueError(f"Unsupported patent file type: {suffix}")


def _read_pdf_text(path: Path) -> str:
    import fitz

    with fitz.open(path) as doc:
        pages = [page.get_text("text") for page in doc]
    text = "\n".join(page for page in pages if page.strip())
    if not text.strip():
        raise ValueError("PDF has no text layer. Please OCR it before importing.")
    return text


def _read_docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def _read_plain_text(path: Path) -> str:
    for encoding in ["utf-8", "utf-8-sig", "gb18030"]:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Text file encoding is not supported. Please save it as UTF-8.")


def _read_xml_text(path: Path) -> str:
    raw = _read_plain_text(path)
    try:
        root = ET.fromstring(raw)
    except ET.ParseError as exc:
        raise ValueError(f"XML patent file is not well formed: {exc}") from exc
    lines = [part.strip() for part in root.itertext() if part and part.strip()]
    text = "\n".join(lines)
    if not text.strip():
        raise ValueError("XML patent file contains no extractable text.")
    return text


def split_patent_sections(text: str) -> list[PatentSection]:
    sections: list[PatentSection] = []
    current_heading: str | None = None
    current_type = SectionType.OTHER
    current_lines: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading_type = _match_heading(line)
        if heading_type:
            if current_heading is not None or current_lines:
                sections.append(
                    PatentSection(
                        type=current_type,
                        heading=current_heading or "未分节内容",
                        text="\n".join(current_lines).strip(),
                        ordinal=len(sections) + 1,
                    )
                )
            current_heading = _canonical_heading(line)
            current_type = heading_type
            current_lines = []
        else:
            current_lines.append(line)

    if current_heading is not None or current_lines:
        sections.append(
            PatentSection(
                type=current_type,
                heading=current_heading or "未分节内容",
                text="\n".join(current_lines).strip(),
                ordinal=len(sections) + 1,
            )
        )
    return sections


def extract_claims(claims_text: str) -> list[PatentClaim]:
    matches = list(re.finditer(r"(?m)^\s*(\d+)[.．、]\s*", claims_text))
    claims: list[PatentClaim] = []
    for index, match in enumerate(matches):
        number = int(match.group(1))
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(claims_text)
        text = claims_text[start:end].strip()
        references = [int(value) for value in re.findall(r"权利要求\s*(\d+)", text)]
        kind = "dependent" if references else "independent"
        claims.append(
            PatentClaim(
                number=number,
                text=text,
                kind=kind,
                references=references,
                category=_classify_claim_category(text),
            )
        )
    return claims


def make_patent_document(document_id: str, source_name: str, text: str) -> PatentDocument:
    sections = split_patent_sections(text)
    title = _infer_title(source_name, text)
    return PatentDocument(
        id=document_id,
        title=title,
        source_name=source_name,
        text=text,
        sections=sections,
        metadata=_infer_metadata(text),
    )


def chunk_document(document: PatentDocument, max_chars: int = 1200) -> list:
    chunks = []
    for section in document.sections:
        pieces = _split_long_text(section.text, max_chars=max_chars) or [""]
        for piece_index, piece in enumerate(pieces, start=1):
            if not piece.strip():
                continue
            chunks.append(
                {
                    "id": f"{document.id}:{section.ordinal}:{piece_index}",
                    "document_id": document.id,
                    "section_type": section.type,
                    "text": piece,
                    "ordinal": len(chunks) + 1,
                    "metadata": {
                        "title": document.title,
                        "source_name": document.source_name,
                        "heading": section.heading,
                        **document.metadata,
                    },
                }
            )
    return chunks


def _match_heading(line: str) -> SectionType | None:
    normalized = _canonical_heading(line)
    return HEADING_MAP.get(normalized)


def _canonical_heading(line: str) -> str:
    line = re.sub(r"^[一二三四五六七八九十\d]+[、.．]\s*", "", line.strip())
    line = re.sub(r"[:：]$", "", line)
    return line


def _infer_title(source_name: str, text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and _match_heading(stripped) is None and len(stripped) <= 80:
            return stripped
    return Path(source_name).stem


def _infer_metadata(text: str) -> dict[str, str]:
    metadata: dict[str, str] = {}
    patterns = {
        "patent_number": r"(?:授权公告号|公开号|公开公告号)[:：]?\s*([A-Z]{0,3}\d+[A-Z]?)",
        "application_number": r"(?:申请号)[:：]?\s*([0-9.]+)",
        "ipc": r"(?:IPC分类号|国际专利分类号)[:：]?\s*([A-Z]\d{2}[A-Z]\s*\d+/\d+)",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            metadata[key] = match.group(1).strip()
    return metadata


def _classify_claim_category(text: str) -> str:
    if re.search(r"(计算机可读)?存储介质|介质", text):
        return "medium"
    if "系统" in text:
        return "system"
    if re.search(r"装置|设备|终端|服务器|处理器", text):
        return "device"
    if "方法" in text:
        return "method"
    return "other"


def _split_long_text(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text] if text.strip() else []
    paragraphs = [paragraph.strip() for paragraph in text.split("\n") if paragraph.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for paragraph in paragraphs:
        if current and current_len + len(paragraph) > max_chars:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(paragraph)
        current_len += len(paragraph)
    if current:
        chunks.append("\n".join(current))
    return chunks
