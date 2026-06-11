from __future__ import annotations

import hashlib
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document

from backend.app.schemas import (
    DeliberationLogEntry,
    DraftPackage,
    ExternalDraftIntakeRun,
    ExternalDraftReviewBundle,
    ExternalDraftSource,
    IntakeIssue,
    SectionConfidence,
    SectionConfidenceItem,
)


PARSER_VERSION = "external-draft-parser-v1"

SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "title": ("发明名称", "题名"),
    "abstract": ("摘要", "摘要附图"),
    "claims": ("权利要求书", "权利要求", "权利要求书正文"),
    "description": ("说明书", "技术领域", "背景技术", "发明内容", "具体实施方式", "实施例"),
    "drawing_description": ("附图说明", "图面说明", "附图简要说明"),
}

DESCRIPTION_SUBHEADINGS = {"技术领域", "背景技术", "发明内容", "具体实施方式", "实施例"}
REQUIRED_SECTIONS = ("claims", "description")
SECTION_LABELS = {
    "title": "发明名称",
    "abstract": "摘要",
    "claims": "权利要求书",
    "description": "说明书",
    "drawing_description": "附图说明",
}


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def working_draft_hash(package: DraftPackage) -> str:
    return hashlib.sha256(package.model_dump_json().encode("utf-8")).hexdigest()


def review_bundle_hash(bundle: ExternalDraftReviewBundle) -> str:
    canonical = bundle.model_copy(update={"report_hash": ""})
    return hashlib.sha256(canonical.model_dump_json().encode("utf-8")).hexdigest()


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_text("\n".join(parts))


def create_external_draft_source(
    *,
    project_id: str,
    source_type: str,
    text: str,
    file_name: str = "",
    raw_path: str = "",
    metadata: dict[str, Any] | None = None,
) -> ExternalDraftSource:
    normalized = normalize_text(text)
    return ExternalDraftSource(
        id=uuid.uuid4().hex,
        project_id=project_id,
        source_type=source_type,
        file_name=file_name or "external-draft.txt",
        content_hash=content_hash(normalized),
        raw_text=normalized,
        raw_path=raw_path,
        metadata=metadata or {},
        created_at=utc_now_iso(),
    )


def parse_external_draft_source(*, project_id: str, source: ExternalDraftSource) -> ExternalDraftIntakeRun:
    logs = [
        DeliberationLogEntry(
            level="info",
            phase="external_draft_intake",
            provider_id="system",
            message="external draft intake started",
            detail=f"source_id={source.id}; source_type={source.source_type}",
        )
    ]
    sections, duplicate_sections, unassigned = parse_sections(source.raw_text)
    issues = intake_issues_from_sections(sections, duplicate_sections, source.raw_text)
    package = package_from_sections(sections)
    status = "needs_review" if any(issue.blocks_quality_run for issue in issues) else "completed"
    return ExternalDraftIntakeRun(
        id=uuid.uuid4().hex,
        project_id=project_id,
        source_id=source.id,
        status=status,
        parser_version=PARSER_VERSION,
        source_hash=source.content_hash,
        parsed_package=package,
        section_confidence=section_confidence_from_sections(sections),
        intake_issues=issues,
        unassigned_fragments=unassigned,
        working_draft_hash=working_draft_hash(package),
        logs=logs,
        created_at=utc_now_iso(),
    )


def normalize_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def parse_sections(text: str) -> tuple[dict[str, str], set[str], list[str]]:
    lines = normalize_text(text).splitlines()
    sections: dict[str, list[str]] = {section: [] for section in SECTION_ALIASES}
    duplicate_sections: set[str] = set()
    unassigned: list[str] = []
    current = ""
    seen: set[str] = set()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if current:
                sections[current].append("")
            continue

        heading = detect_heading(line)
        if heading:
            heading_text = strip_heading_marker(line).rstrip("：:")
            if current == "description" and heading == "description" and _is_description_subheading(heading_text):
                sections[current].append(raw_line)
                continue

            current = heading
            if heading in seen:
                duplicate_sections.add(heading)
            seen.add(heading)
            if heading == "title" and _canonical_heading_text(heading_text) not in _canonical_aliases("title"):
                sections["title"].append(heading_text)
            continue

        inline_heading = _inline_heading_content(line)
        if inline_heading:
            heading, inline_content = inline_heading
            current = heading
            if heading in seen:
                duplicate_sections.add(heading)
            seen.add(heading)
            sections[current].append(inline_content)
            continue

        markdown_title = _markdown_title(line)
        if markdown_title and not sections["title"]:
            sections["title"].append(markdown_title)
            continue

        if current:
            sections[current].append(raw_line)
        else:
            unassigned.append(raw_line)

    compacted = {section: normalize_text("\n".join(value)) for section, value in sections.items()}
    return compacted, duplicate_sections, [fragment for fragment in unassigned if fragment.strip()]


def detect_heading(line: str) -> str:
    cleaned = _canonical_heading_text(line)
    for section in SECTION_ALIASES:
        if cleaned in _canonical_aliases(section):
            return section
    return ""


def strip_heading_marker(line: str) -> str:
    stripped = re.sub(r"^#{1,6}\s*", "", line.strip())
    stripped = re.sub(r"^\*\*(.+)\*\*$", r"\1", stripped)
    return stripped.strip()


def _canonical_heading_text(line: str) -> str:
    stripped = strip_heading_marker(line)
    stripped = stripped.strip(" \t\r\n:：;；。．.、-—_[]【】()（）<>《》「」『』“”\"'")
    return re.sub(r"\s+", "", stripped)


def _canonical_aliases(section: str) -> set[str]:
    return {_canonical_heading_text(alias) for alias in SECTION_ALIASES[section]}


def _is_description_subheading(heading_text: str) -> bool:
    return _canonical_heading_text(heading_text) in {_canonical_heading_text(text) for text in DESCRIPTION_SUBHEADINGS}


def _inline_heading_content(line: str) -> tuple[str, str] | None:
    stripped = strip_heading_marker(line)
    for delimiter in ("：", ":"):
        heading_text, separator, inline_content = stripped.partition(delimiter)
        if not separator:
            continue
        heading_key = _canonical_heading_text(heading_text)
        for section in SECTION_ALIASES:
            if heading_key in _canonical_aliases(section):
                content = inline_content.strip()
                if content:
                    return section, content
    return None


def package_from_sections(sections: dict[str, str]) -> DraftPackage:
    return DraftPackage(
        title=sections.get("title") or "未命名发明",
        abstract=sections.get("abstract", ""),
        claims=sections.get("claims", ""),
        description=sections.get("description", ""),
        drawing_description=sections.get("drawing_description", ""),
        mermaid="",
        image_prompt="",
        review_findings=[],
        citations=[],
        generation_logs=["external_draft_intake: parsed from external source"],
    )


def section_confidence_from_sections(sections: dict[str, str]) -> SectionConfidence:
    return SectionConfidence(
        title=confidence_item(sections, "title"),
        abstract=confidence_item(sections, "abstract"),
        claims=confidence_item(sections, "claims"),
        description=confidence_item(sections, "description"),
        drawing_description=confidence_item(sections, "drawing_description"),
    )


def confidence_item(sections: dict[str, str], section: str) -> SectionConfidenceItem:
    if not sections.get(section, "").strip():
        return SectionConfidenceItem(
            score=0.0,
            source_markers=[],
            warnings=[f"未识别{SECTION_LABELS.get(section, section)}章节"],
        )
    score = 0.95 if section in REQUIRED_SECTIONS else 0.85
    return SectionConfidenceItem(score=score, source_markers=list(SECTION_ALIASES[section]), warnings=[])


def intake_issues_from_sections(
    sections: dict[str, str], duplicate_sections: Iterable[str], raw_text: str
) -> list[IntakeIssue]:
    issues: list[IntakeIssue] = []
    for section in REQUIRED_SECTIONS:
        if not sections.get(section, "").strip():
            issues.append(
                IntakeIssue(
                    id=f"intake-missing-{section}",
                    category="missing_section",
                    severity="high",
                    section=section,
                    message=f"未识别{SECTION_LABELS[section]}章节。",
                    suggested_action="在章节确认界面补充该章节后再运行质量检查。",
                    blocks_quality_run=True,
                )
            )

    for section in sorted(duplicate_sections):
        issues.append(
            IntakeIssue(
                id=f"intake-duplicate-{section}",
                category="duplicate_section",
                severity="medium",
                section=section,
                message=f"检测到重复的{SECTION_LABELS.get(section, section)}章节标题。",
                suggested_action="确认重复章节是否应合并为同一章节。",
                blocks_quality_run=False,
            )
        )

    claims = sections.get("claims", "")
    if claims.strip() and not _has_standard_claim_numbering(claims):
        issues.append(
            IntakeIssue(
                id="intake-malformed-claim-numbering",
                category="malformed_claim_numbering",
                severity="medium",
                section="claims",
                message="权利要求书未检测到标准的权利要求1编号。",
                suggested_action="将第一项权利要求改为“1.”、“1、”或“权利要求1”开头。",
                blocks_quality_run=False,
            )
        )

    if re.search(r"(?i)(prompt|generation_logs|attorney_memo|system_trace)", raw_text):
        issues.append(
            IntakeIssue(
                id="intake-suspected-internal-text",
                category="suspected_internal_text",
                severity="medium",
                section="raw_text",
                message="外部稿中疑似包含内部过程文本。",
                suggested_action="运行正式稿编译前确认该类文本不会进入正式提交稿。",
                blocks_quality_run=False,
            )
        )
    return issues


def external_draft_review_bundle_to_markdown(bundle: ExternalDraftReviewBundle) -> str:
    lines = [
        "# EXTERNAL_DRAFT_REVIEW_BUNDLE",
        "",
        f"- project_id: {bundle.project_id}",
        f"- source_id: {bundle.source_id or '无'}",
        f"- intake_run_id: {bundle.intake_run_id or '无'}",
        f"- initial_score: {bundle.initial_score if bundle.initial_score is not None else '无'}",
        f"- latest_score: {bundle.latest_score if bundle.latest_score is not None else '无'}",
        f"- accepted_patch_ids: {', '.join(bundle.accepted_patch_ids) or '无'}",
        f"- completion_run_ids: {', '.join(bundle.completion_run_ids) or '无'}",
        f"- official_compile_run_id: {bundle.official_compile_run_id or '无'}",
        f"- post_draft_review_run_id: {bundle.post_draft_review_run_id or '无'}",
        f"- export_allowed: {str(bundle.export_allowed).lower()}",
        f"- report_hash: {bundle.report_hash or '无'}",
        "",
        "本报告是内部提质侧车文件，不进入正式申请正文。",
    ]
    return "\n".join(lines).strip() + "\n"


def _markdown_title(line: str) -> str:
    match = re.match(r"^#\s+(.+)$", line)
    if not match:
        return ""
    return match.group(1).strip()


def _has_standard_claim_numbering(claims: str) -> bool:
    return bool(re.search(r"(?m)^\s*(?:1[\.、．]|权利要求\s*1(?:\D|$))", claims))
