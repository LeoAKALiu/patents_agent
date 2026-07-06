from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.app.schemas import DraftPackage


def export_docx(package: DraftPackage, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(package.title, level=0)
    _add_section(doc, "摘要", package.abstract)
    _add_section(doc, "权利要求书", package.claims)
    _add_section(doc, "说明书", package.description)
    _add_section(doc, "附图说明", package.drawing_description)
    doc.save(output_path)
    return output_path


def package_to_markdown(package: DraftPackage) -> str:
    return f"""# {package.title}

## 摘要
{package.abstract}

## 权利要求书
{package.claims}

## 说明书
{package.description}

## 附图说明
{package.drawing_description}
"""


def _add_section(doc: Document, heading: str, text: str) -> None:
    doc.add_heading(heading, level=1)
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)
