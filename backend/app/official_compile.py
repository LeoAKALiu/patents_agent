from __future__ import annotations

import csv
import hashlib
import io
import re
import uuid
from pathlib import Path

from docx import Document

from backend.app.schemas import (
    DeliberationLogEntry,
    DraftPackage,
    OfficialCompileRun,
    OfficialDraftPackage,
    OfficialFigurePlanItem,
)


CROSS_PROJECT_TITLE = "基于边缘端动态推理的无人机飞行中任务调整方法"
REQUIRED_SECTIONS = ("abstract", "claims", "description", "drawing_description")
HARD_GATED_SECTIONS = ("title", *REQUIRED_SECTIONS)
RESIDUAL_INTERNAL_PATTERNS = (
    "support_gap",
    "support_gaps",
    "evidence_id",
    "evidence_refs",
    "research_ledger",
    "revision_ledger",
    "source_ledger",
    "publication_number",
    "patent_url",
    "source_url",
    "source_id",
    "source_label",
    "material_id",
    "证据编号",
    "材料编号",
    "来源标签",
    "引用来源",
    "引用链接",
    "证据来源",
    "evidence_status",
    "verification_status",
    "internal_only",
    "修订记录",
    "patents.google.com",
    "generation_logs",
    "image_prompt",
    "prompt",
    "diagram",
    "attorney_memo",
    "system_trace",
    "official_safe_patches",
    "好的，下面",
    "好的，根据",
    "待验证",
    "主席修订",
    "需在提交前补充",
    "提交前补充",
    "方法方法",
    "颠覆",
)
AUTO_CLEANED_TEXT_CATEGORY = "auto_cleaned_text"
INTERNAL_FIELD_RE = re.compile(
    r"""^\s*["']?(image_prompt|prompt|diagram|generation_logs|attorney_memo|system_trace|official_safe_patches|revision_ledger|source_ledger|修订记录)["']?\s*[:：=]""",
    re.IGNORECASE | re.MULTILINE,
)
EVIDENCE_METADATA_FIELD_RE = re.compile(
    r"""^\s*["']?(evidence_id|evidence_refs|research_ledger|publication_number|patent_url|source_url|url|material_id|source_id|source_label|sources?|references?|citations?|materials?|patent_point|证据编号|材料编号|来源标签|引用来源|引用链接|证据来源|参考资料|参考文献|资料来源|依据材料|支撑材料)["']?\s*[:：=]""",
    re.IGNORECASE,
)
EMPTY_OFFICIAL_FIELD_RE = re.compile(
    r"""^\s*["']?(title|abstract|claims|description|drawing_description)["']?\s*[:：=]\s*["']?\s*["']?\s*,?\s*$""",
    re.IGNORECASE,
)
OFFICIAL_FIELD_WRAPPER_RE = re.compile(
    r"""^\s*[{\[,]?\s*["']?(title|abstract|claims|description|drawing_description)["']?\s*[:：=]\s*.+""",
    re.IGNORECASE,
)
JSON_WRAPPER_RE = re.compile(r"^[{}\[\],]+$")
URL_RE = re.compile(r"https?://", re.IGNORECASE)
BRACKETED_EVIDENCE_CITATION_RE = re.compile(
    r"""(?:\[[^\]]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^\]]+\]|【[^】]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^】]+】|\[(?:EV|EVIDENCE)-[A-Z0-9_-]+\]|【(?:EV|EVIDENCE)-[A-Z0-9_-]+】)""",
    re.IGNORECASE,
)
PARENTHETICAL_EVIDENCE_CITATION_RE = re.compile(
    r"""(?:\([^)]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^)]+\)|（[^）]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^）]+）)""",
    re.IGNORECASE,
)
XML_EVIDENCE_TAG_RE = re.compile(
    r"""<\s*/?\s*(?:evidence|source|citation|ref|reference|references|material|materials)\b[^>]*>""",
    re.IGNORECASE,
)
HTML_COMMENT_EVIDENCE_CITATION_RE = re.compile(
    r"""<!--[^>]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^>]*-->""",
    re.IGNORECASE,
)
HTML_ATTRIBUTE_EVIDENCE_CITATION_RE = re.compile(
    r"""<[^>]+\s(?:data-)?(?:evidence(?:[-_]?(?:id|ref|refs))?|source(?:[-_]?(?:id|label|ref))?|citation|ref|material(?:[-_]?id)?|证据|来源|引用|材料)\s*=\s*["'][^"']+["'][^>]*>""",
    re.IGNORECASE,
)
HTML_CLASS_ID_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*[a-z][\w:-]*\b(?=[^>]*\b(?:class|id)\s*=\s*(?:"[^"]*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[-_\s:：=]|$)|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^"]*"|'[^']*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[-_\s:：=]|$)|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^']*'))[^>]*>""",
    re.IGNORECASE,
)
HTML_DATA_VALUE_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*[a-z][\w:-]*\b(?=[^>]*\bdata-[\w:-]+\s*=\s*(?:"[^"]*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^"]*"|'[^']*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^']*'))[^>]*>""",
    re.IGNORECASE,
)
HTML_EVENT_HANDLER_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*[a-z][\w:-]*\b(?=[^>]*\bon[a-z][\w:-]*\s*=\s*(?:"[^"]*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^"]*"|'[^']*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^']*'))[^>]*>""",
    re.IGNORECASE,
)
HTML_URL_ATTRIBUTE_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*[a-z][\w:-]*\b(?=[^>]*\b(?:href|src|action|poster|srcdoc)\s*=\s*(?:"[^"]*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^"]*"|'[^']*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^']*'))[^>]*>""",
    re.IGNORECASE,
)
HTML_SRCSET_ATTRIBUTE_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*[a-z][\w:-]*\b(?=[^>]*\bsrcset\s*=\s*(?:"[^"]*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^"]*"|'[^']*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^']*'))[^>]*>""",
    re.IGNORECASE,
)
HTML_IMAGE_ATTRIBUTE_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*img\b(?=[^>]*\b(?:alt|title|aria-label)\s*=\s*["'][^"']*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^"']*["'])[^>]*>""",
    re.IGNORECASE,
)
HTML_ACCESSIBLE_ATTRIBUTE_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?!img\b)[a-z][\w:-]*\b(?=[^>]*\b(?:alt|title|aria-label)\s*=\s*["'][^"']*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^"']*["'])[^>]*>""",
    re.IGNORECASE,
)
HTML_META_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*meta\b(?=[^>]*\b(?:name|property|itemprop)\s*=\s*["'][^"']*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)[^"']*["'])(?=[^>]*\bcontent\s*=\s*["'][^"']+["'])[^>]*>""",
    re.IGNORECASE,
)
HTML_JSON_LD_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*script\b(?=[^>]*\btype\s*=\s*["']application/ld\+json["'])[^>]*>.*?["'](?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[_-]?(?:id|refs?|label|编号|标签|来源))?["']\s*:""",
    re.IGNORECASE | re.DOTALL,
)
HTML_JSON_SCRIPT_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*script\b(?=[^>]*\btype\s*=\s*["']application/(?:json|x-json)(?:\s*;[^"']*)?["'])[^>]*>.*?["'](?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[_-]?(?:id|refs?|label|编号|标签|来源))?["']\s*:""",
    re.IGNORECASE | re.DOTALL,
)
HTML_SCRIPT_TEXT_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*script\b[^>]*>.*?(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+).*?<\s*/\s*script\s*>""",
    re.IGNORECASE | re.DOTALL,
)
MARKDOWN_FOOTNOTE_EVIDENCE_CITATION_RE = re.compile(
    r"""(?:^|\n)\s*\[\^[^\]\n]+\]\s*:\s*[^\n]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]?[^\n]*""",
    re.IGNORECASE,
)
MARKDOWN_REFERENCE_EVIDENCE_CITATION_RE = re.compile(
    r"""(?:^|\n)\s*\[(?:evidence|source|citation|ref|material|证据|来源|引用|材料)[^\]\n]*\]\s*:\s*[^\n]+""",
    re.IGNORECASE,
)
MARKDOWN_TABLE_EVIDENCE_CITATION_RE = re.compile(
    r"""^\s*\|(?:[^|\n]*\|)*\s*(?:evidence(?:_id|_refs)?|source(?:_id|_label)?|citation|ref|material(?:_id)?|证据(?:编号|来源)?|来源(?:标签)?|引用(?:来源)?|材料(?:编号)?)\s*(?:[:：=][^|\n]*)?\|""",
    re.IGNORECASE | re.MULTILINE,
)
MARKDOWN_LIST_EVIDENCE_CITATION_RE = re.compile(
    r"""^\s*(?:[-*+]|\d+[.)])\s+(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[_-]?(?:id|refs?|label|编号|标签|来源))?\s*[:：=]\s*\S""",
    re.IGNORECASE | re.MULTILINE,
)
MARKDOWN_BLOCKQUOTE_EVIDENCE_CITATION_RE = re.compile(
    r"""^\s*>+\s*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[_-]?(?:id|refs?|label|编号|标签|来源))?\s*[:：=]\s*\S""",
    re.IGNORECASE | re.MULTILINE,
)
MARKDOWN_IMAGE_ALT_EVIDENCE_CITATION_RE = re.compile(
    r"""!\[[^\]\n]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^\]\n]+\]\([^\)\n]*\)""",
    re.IGNORECASE,
)
MARKDOWN_LINK_TITLE_EVIDENCE_CITATION_RE = re.compile(
    r"""!?\[[^\]\n]*\]\([^\)\n]*(?:"[^"\n]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^"\n]*"|'[^'\n]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^'\n]*')[^\)\n]*\)""",
    re.IGNORECASE,
)
HTML_VISIBLE_TEXT_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?P<html_visible_text_tag>span|sup|sub|small|em|i|b|strong|mark|p|div)\b[^>]*>[^<]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^<]+<\s*/\s*(?P=html_visible_text_tag)\s*>""",
    re.IGNORECASE,
)
HTML_HIDDEN_TEXT_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?P<html_hidden_text_tag>template|noscript)\b[^>]*>.*?(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=].*?<\s*/\s*(?P=html_hidden_text_tag)\s*>""",
    re.IGNORECASE | re.DOTALL,
)
HTML_CAPTION_TEXT_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?P<html_caption_text_tag>figcaption|caption)\b[^>]*>[^<]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^<]+<\s*/\s*(?P=html_caption_text_tag)\s*>""",
    re.IGNORECASE,
)
SVG_TITLE_DESC_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?P<svg_title_desc_tag>title|desc)\b[^>]*>[^<]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^<]+<\s*/\s*(?P=svg_title_desc_tag)\s*>""",
    re.IGNORECASE,
)
SVG_TEXT_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?P<svg_text_tag>text|tspan)\b[^>]*>[^<]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^<]+<\s*/\s*(?P=svg_text_tag)\s*>""",
    re.IGNORECASE,
)
HTML_STYLE_TAG_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*style\b[^>]*>.*?(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=].*?<\s*/\s*style\s*>""",
    re.IGNORECASE | re.DOTALL,
)
HTML_INLINE_STYLE_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?!style\b)[a-z][\w:-]*\b(?=[^>]*\bstyle\s*=\s*(?:"[^"]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^"]*"|'[^']*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^']*'))[^>]*>""",
    re.IGNORECASE,
)
HTML_FORM_FIELD_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?:input|textarea|select|option)\b(?=[^>]*(?:\b(?:value|data-value)\s*=\s*(?:"[^"]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^"]*"|'[^']*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^']*')|\b(?:name|id)\s*=\s*(?:"[^"]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)[^"]*"|'[^']*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)[^']*')(?=[^>]*\b(?:value|data-value)\s*=\s*(?:"[^"]+"|'[^']+"))))[^>]*>""",
    re.IGNORECASE,
)
HTML_SEMANTIC_METADATA_EVIDENCE_CITATION_RE = re.compile(
    r"""<\s*(?!meta\b)[a-z][\w:-]*\b(?=[^>]*\b(?:itemprop|property)\s*=\s*(?:"[^"]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)[^"]*"|'[^']*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)[^']*'))(?=[^>]*\b(?:content|value|data-content)\s*=\s*(?:"[^"]+"|'[^']+"))[^>]*>""",
    re.IGNORECASE,
)
HTML_ENTITY_EVIDENCE_CITATION_RE = re.compile(
    r"""&lt;\s*/?\s*(?:evidence|source|citation|ref|reference|references|material|materials)\b[^\n]*?&gt;""",
    re.IGNORECASE,
)
YAML_FRONT_MATTER_EVIDENCE_KEY_RE = re.compile(
    r"""^\s*(?:evidence|证据)\s*[:：=]""",
    re.IGNORECASE,
)
TOML_FRONT_MATTER_EVIDENCE_KEY_RE = re.compile(
    r"""^\s*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[_-]?(?:id|refs?|label|编号|标签|来源))?\s*=""",
    re.IGNORECASE,
)
INI_EVIDENCE_SECTION_RE = re.compile(
    r"""^\s*\[\s*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[_\-\s]*(?:id|refs?|label|编号|标签|来源))?\s*\]\s*$""",
    re.IGNORECASE,
)
INI_EVIDENCE_KEY_RE = re.compile(
    r"""^\s*(?:id|refs?|label|url|source|evidence|citation|ref|material|编号|标签|来源|证据|引用|材料)\s*=""",
    re.IGNORECASE,
)
CSV_EVIDENCE_METADATA_HEADER_RE = re.compile(
    r"""^(?:evidence(?:_id|_refs?)?|source(?:_id|_label|_ref)?|citation|ref|material(?:_id)?|证据(?:编号|来源)?|来源(?:标签)?|引用(?:来源)?|材料(?:编号)?)$""",
    re.IGNORECASE,
)
MARKDOWN_CODE_FENCE_RE = re.compile(r"""```(?P<language>[A-Za-z0-9_-]*)[^\n]*\n(?P<body>.*?)\n```""", re.DOTALL)
JSON_EVIDENCE_METADATA_KEY_RE = re.compile(
    r""""(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[_-]?(?:id|refs?|label|编号|标签|来源))?"\s*:""",
    re.IGNORECASE,
)
ASCIIDOC_ATTRIBUTE_EVIDENCE_CITATION_RE = re.compile(
    r"""^\s*:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[-_\s]?(?:id|refs?|label|编号|标签|来源))?\s*:\s*\S""",
    re.IGNORECASE | re.MULTILINE,
)
LATEX_COMMAND_EVIDENCE_CITATION_RE = re.compile(
    r"""\\(?:cite\w*|footnote|thanks|ref|autoref)\s*(?:\[[^\]]*(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=][^\]]*\]\s*)?\{[^{}]*(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[:：=]|(?:EV|EVIDENCE)-[A-Z0-9_-]+)[^{}]*\}""",
    re.IGNORECASE,
)
BIBTEX_ENTRY_EVIDENCE_CITATION_RE = re.compile(
    r"""@\w+\s*\{(?=[\s\S]*?(?:(?:evidence|source|citation|ref|material|证据|来源|引用|材料)\s*[=:：]|(?:EV|EVIDENCE)-[A-Z0-9_-]+))[\s\S]*?\n\s*\}""",
    re.IGNORECASE,
)
RST_DIRECTIVE_EVIDENCE_CITATION_RE = re.compile(
    r"""^\s*\.\.\s+(?:evidence|source|citation|ref|material|证据|来源|引用|材料)(?:[-_\s]?(?:id|refs?|label|编号|标签|来源))?\s*::\s*\S""",
    re.IGNORECASE | re.MULTILINE,
)


def source_draft_hash(package: DraftPackage) -> str:
    return hashlib.sha256(package.model_dump_json().encode("utf-8")).hexdigest()


def official_package_hash(package: OfficialDraftPackage) -> str:
    canonical = package.model_copy(update={"official_package_hash": ""})
    return hashlib.sha256(canonical.model_dump_json().encode("utf-8")).hexdigest()


def clean_source_draft_for_official_compile(
    package: DraftPackage,
) -> tuple[DraftPackage, list[dict[str, str]], list[dict[str, str]]]:
    contamination_removed: list[dict[str, str]] = []
    sidecar_notes: list[dict[str, str]] = []
    cleaned_fields = {
        section: _clean_section(
            section=section,
            text=getattr(package, section),
            contamination_removed=contamination_removed,
            sidecar_notes=sidecar_notes,
        )
        for section in ("title", *REQUIRED_SECTIONS)
    }
    return package.model_copy(update=cleaned_fields), contamination_removed, sidecar_notes


class OfficialDraftCompiler:
    def compile(self, project_id: str, package: DraftPackage) -> OfficialCompileRun:
        run_id = uuid.uuid4().hex
        now = _utc_now_iso()
        package_hash = source_draft_hash(package)
        blocked_items: list[dict[str, str]] = []
        logs: list[DeliberationLogEntry] = [
            DeliberationLogEntry(
                level="info",
                phase="official_compile",
                message="official draft compile started",
                detail=f"source_draft_hash={package_hash}",
            )
        ]

        source_text = "\n".join(
            [
                package.title,
                package.abstract,
                package.claims,
                package.description,
                package.drawing_description,
            ]
        )
        if CROSS_PROJECT_TITLE in source_text:
            blocked_items.append(
                {
                    "category": "cross_project_contamination",
                    "section": "draft_package",
                    "pattern": CROSS_PROJECT_TITLE,
                    "message": "Detected title from another project in draft text.",
                }
            )

        for section in HARD_GATED_SECTIONS:
            source_section_text = getattr(package, section)
            internal_field = INTERNAL_FIELD_RE.search(source_section_text)
            if internal_field:
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": internal_field.group(1).lower(),
                        "message": "Draft text contains internal field metadata that must not appear in official text.",
                    }
                )
            if _contains_fenced_json_evidence_metadata(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "fenced_json_metadata_citation",
                        "message": "Draft text contains fenced JSON evidence metadata that must not appear in official text.",
                    }
                )
            if ASCIIDOC_ATTRIBUTE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "asciidoc_attribute_citation",
                        "message": "Draft text contains AsciiDoc evidence metadata that must not appear in official text.",
                    }
                )
            if LATEX_COMMAND_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "latex_command_citation",
                        "message": "Draft text contains LaTeX evidence metadata that must not appear in official text.",
                    }
                )
            if BIBTEX_ENTRY_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "bibtex_entry_citation",
                        "message": "Draft text contains BibTeX evidence metadata that must not appear in official text.",
                    }
                )
            if RST_DIRECTIVE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "rst_directive_citation",
                        "message": "Draft text contains reStructuredText evidence metadata that must not appear in official text.",
                    }
                )
            if MARKDOWN_LIST_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_list_citation",
                        "message": "Draft text contains Markdown list evidence metadata that must not appear in official text.",
                    }
                )
            if MARKDOWN_BLOCKQUOTE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_blockquote_citation",
                        "message": "Draft text contains Markdown blockquote evidence metadata that must not appear in official text.",
                    }
                )
            if MARKDOWN_IMAGE_ALT_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_image_alt_citation",
                        "message": "Draft text contains Markdown image alt evidence metadata that must not appear in official text.",
                    }
                )
            if MARKDOWN_LINK_TITLE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_link_title_citation",
                        "message": "Draft text contains Markdown link title evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_JSON_SCRIPT_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_json_script_citation",
                        "message": "Draft text contains HTML JSON script evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_SCRIPT_TEXT_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_script_text_citation",
                        "message": "Draft text contains HTML script evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_CLASS_ID_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_class_id_citation",
                        "message": "Draft text contains HTML class/id evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_DATA_VALUE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_data_value_citation",
                        "message": "Draft text contains HTML data attribute evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_EVENT_HANDLER_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_event_handler_citation",
                        "message": "Draft text contains HTML event handler evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_URL_ATTRIBUTE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_url_attribute_citation",
                        "message": "Draft text contains HTML URL attribute evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_SRCSET_ATTRIBUTE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_srcset_attribute_citation",
                        "message": "Draft text contains HTML srcset attribute evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_IMAGE_ATTRIBUTE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_image_attribute_citation",
                        "message": "Draft text contains HTML image attribute evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_ACCESSIBLE_ATTRIBUTE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_accessible_attribute_citation",
                        "message": "Draft text contains HTML accessible attribute evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_VISIBLE_TEXT_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_visible_text_citation",
                        "message": "Draft text contains HTML visible evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_HIDDEN_TEXT_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_hidden_text_citation",
                        "message": "Draft text contains HTML hidden evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_CAPTION_TEXT_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_caption_text_citation",
                        "message": "Draft text contains HTML caption evidence metadata that must not appear in official text.",
                    }
                )
            if SVG_TITLE_DESC_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "svg_title_desc_citation",
                        "message": "Draft text contains SVG title/desc evidence metadata that must not appear in official text.",
                    }
                )
            if SVG_TEXT_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "svg_text_citation",
                        "message": "Draft text contains SVG visible text evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_STYLE_TAG_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_style_tag_citation",
                        "message": "Draft text contains HTML style evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_INLINE_STYLE_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_inline_style_citation",
                        "message": "Draft text contains HTML inline style evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_FORM_FIELD_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_form_field_citation",
                        "message": "Draft text contains HTML form-field evidence metadata that must not appear in official text.",
                    }
                )
            if HTML_SEMANTIC_METADATA_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_semantic_metadata_citation",
                        "message": "Draft text contains HTML semantic metadata evidence that must not appear in official text.",
                    }
                )
            if HTML_ENTITY_EVIDENCE_CITATION_RE.search(source_section_text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_entity_citation",
                        "message": "Draft text contains HTML entity-escaped evidence metadata that must not appear in official text.",
                    }
                )

        cleaned_package, contamination_removed, sidecar_notes = clean_source_draft_for_official_compile(package)
        cleaned_title = cleaned_package.title
        cleaned = {section: getattr(cleaned_package, section) for section in REQUIRED_SECTIONS}

        if not cleaned_title.strip():
            blocked_items.append(
                {
                    "category": "empty_required_section",
                    "section": "title",
                    "pattern": "empty_after_cleaning",
                    "message": "title is empty after removing internal contamination.",
                }
            )

        for section, text in cleaned.items():
            if not text.strip():
                blocked_items.append(
                    {
                        "category": "empty_required_section",
                        "section": section,
                        "pattern": "empty_after_cleaning",
                        "message": f"{section} is empty after removing internal contamination.",
                    }
                )

        all_cleaned_text = {"title": cleaned_title, **cleaned}
        for item in contamination_removed:
            if item["section"] not in HARD_GATED_SECTIONS:
                continue
            if item["category"] == AUTO_CLEANED_TEXT_CATEGORY:
                continue
            blocked_items.append(
                {
                    "category": "official_hygiene_contamination",
                    "section": item["section"],
                    "pattern": item["pattern"],
                    "message": "Official draft field contains generator, format, prompt, memo, support-gap, or process-trace text; revise the source draft and recompile.",
                }
            )

        for section, text in all_cleaned_text.items():
            comparable_text = text.lower()
            for pattern in RESIDUAL_INTERNAL_PATTERNS:
                if pattern in comparable_text:
                    blocked_items.append(
                        {
                            "category": "residual_internal_text",
                            "section": section,
                            "pattern": pattern,
                            "message": "Cleaned official text still contains internal drafting text.",
                        }
                    )
            if URL_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "url",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if BRACKETED_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "bracketed_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if PARENTHETICAL_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "parenthetical_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if XML_EVIDENCE_TAG_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "xml_evidence_tag",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_COMMENT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_comment_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_ATTRIBUTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_attribute_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_CLASS_ID_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_class_id_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_DATA_VALUE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_data_value_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_EVENT_HANDLER_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_event_handler_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_URL_ATTRIBUTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_url_attribute_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_SRCSET_ATTRIBUTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_srcset_attribute_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_IMAGE_ATTRIBUTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_image_attribute_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_ACCESSIBLE_ATTRIBUTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_accessible_attribute_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_META_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_meta_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_JSON_LD_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_json_ld_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_JSON_SCRIPT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_json_script_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_SCRIPT_TEXT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_script_text_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if MARKDOWN_FOOTNOTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_footnote_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if MARKDOWN_REFERENCE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_reference_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if MARKDOWN_TABLE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_table_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if MARKDOWN_LIST_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_list_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if MARKDOWN_BLOCKQUOTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_blockquote_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if MARKDOWN_IMAGE_ALT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_image_alt_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if MARKDOWN_LINK_TITLE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "markdown_link_title_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_VISIBLE_TEXT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_visible_text_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_HIDDEN_TEXT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_hidden_text_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_CAPTION_TEXT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_caption_text_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if SVG_TITLE_DESC_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "svg_title_desc_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if SVG_TEXT_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "svg_text_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_STYLE_TAG_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_style_tag_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_INLINE_STYLE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_inline_style_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_FORM_FIELD_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_form_field_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_SEMANTIC_METADATA_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_semantic_metadata_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if HTML_ENTITY_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "html_entity_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if ASCIIDOC_ATTRIBUTE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "asciidoc_attribute_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if LATEX_COMMAND_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "latex_command_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if BIBTEX_ENTRY_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "bibtex_entry_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if RST_DIRECTIVE_EVIDENCE_CITATION_RE.search(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "rst_directive_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if _contains_yaml_front_matter_evidence_metadata(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "yaml_front_matter_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if _contains_toml_front_matter_evidence_metadata(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "toml_front_matter_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if _contains_ini_section_evidence_metadata(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "ini_section_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )
            if _contains_csv_evidence_metadata(text):
                blocked_items.append(
                    {
                        "category": "residual_internal_text",
                        "section": section,
                        "pattern": "csv_metadata_citation",
                        "message": "Cleaned official text still contains internal drafting text.",
                    }
                )

        if blocked_items:
            logs.append(
                DeliberationLogEntry(
                    level="warn",
                    phase="official_compile",
                    message="official draft compile blocked",
                    detail=f"blocked_items={len(blocked_items)}",
                )
            )
            return OfficialCompileRun(
                id=run_id,
                project_id=project_id,
                status="blocked",
                source_draft_hash=package_hash,
                contamination_removed=contamination_removed,
                blocked_items=blocked_items,
                sidecar_notes=sidecar_notes,
                logs=logs,
                created_at=now,
                updated_at=now,
            )

        official_package = OfficialDraftPackage(
            title=cleaned_title,
            abstract=cleaned["abstract"],
            claims=cleaned["claims"],
            description=cleaned["description"],
            drawing_description=cleaned["drawing_description"],
            figure_plan=_parse_figure_plan(cleaned["drawing_description"]),
            compile_warnings=[],
            source_draft_hash=package_hash,
        )
        official_hash = official_package_hash(official_package)
        official_package.official_package_hash = official_hash
        logs.append(
            DeliberationLogEntry(
                level="info",
                phase="official_compile",
                message="official draft compile completed",
                detail=f"official_package_hash={official_hash}",
            )
        )
        return OfficialCompileRun(
            id=run_id,
            project_id=project_id,
            status="completed",
            source_draft_hash=package_hash,
            official_package_hash=official_hash,
            official_package=official_package,
            contamination_removed=contamination_removed,
            blocked_items=[],
            sidecar_notes=sidecar_notes,
            logs=logs,
            created_at=now,
            updated_at=now,
        )


def official_package_to_markdown(package: OfficialDraftPackage) -> str:
    lines = [
        f"# {package.title}",
        "",
        "## 摘要",
        package.abstract,
        "",
        "## 权利要求书",
        package.claims,
        "",
        "## 说明书",
        package.description,
        "",
        "## 附图说明",
        package.drawing_description,
    ]
    if package.figure_plan:
        lines.extend(["", "## 附图计划"])
        for item in package.figure_plan:
            lines.append(f"- {item.figure_no}：{item.title}。{item.description}")
    return "\n".join(lines).strip() + "\n"


def export_official_package_docx(package: OfficialDraftPackage, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(package.title, level=0)
    _add_docx_section(doc, "摘要", package.abstract)
    _add_docx_section(doc, "权利要求书", package.claims)
    _add_docx_section(doc, "说明书", package.description)
    _add_docx_section(doc, "附图说明", package.drawing_description)
    if package.figure_plan:
        doc.add_heading("附图计划", level=1)
        for item in package.figure_plan:
            doc.add_paragraph(f"{item.figure_no}：{item.title}。{item.description}")
    doc.save(output_path)
    return output_path


def official_compile_run_to_markdown(run: OfficialCompileRun) -> str:
    lines = [
        "# OFFICIAL_COMPILE_RUN",
        "",
        f"- run_id: {run.id}",
        f"- project_id: {run.project_id}",
        f"- status: {run.status}",
        f"- source_draft_hash: {run.source_draft_hash}",
        f"- official_package_hash: {run.official_package_hash}",
        "",
        "## Blocked Items",
    ]
    lines.extend(_dict_item_lines(run.blocked_items))
    lines.extend(["", "## Contamination Removed"])
    lines.extend(_dict_item_lines(run.contamination_removed))
    lines.extend(["", "## Sidecar Notes"])
    lines.extend(_dict_item_lines(run.sidecar_notes))
    if run.official_package:
        lines.extend(["", "## Official Package", "", official_package_to_markdown(run.official_package).strip()])
    lines.extend(["", "## Logs"])
    for log in run.logs:
        lines.append(f"- [{log.level}] {log.phase}: {log.message} {log.detail}".strip())
    return "\n".join(lines).strip() + "\n"


def _clean_section(
    *,
    section: str,
    text: str,
    contamination_removed: list[dict[str, str]],
    sidecar_notes: list[dict[str, str]],
) -> str:
    kept: list[str] = []
    in_fence = False
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if kept and kept[-1] != "":
                kept.append("")
            continue
        removal = _removal_for_line(line, in_fence)
        if line.startswith("```"):
            in_fence = not in_fence
        if removal:
            item = {
                "category": removal["category"],
                "section": section,
                "pattern": removal["pattern"],
                "text": line,
            }
            contamination_removed.append(item)
            if removal["category"] == "support_gap":
                sidecar_notes.append(item.copy())
            continue
        cleaned_line = _strip_inline_markdown(line)
        cleaned_line = _clean_inline_official_text(
            section=section,
            text=cleaned_line,
            contamination_removed=contamination_removed,
        )
        if cleaned_line:
            kept.append(cleaned_line)
    return "\n".join(kept).strip()


def _clean_inline_official_text(
    *,
    section: str,
    text: str,
    contamination_removed: list[dict[str, str]],
) -> str:
    cleaned = text
    replacements = (
        ("方法方法", "方法"),
        ("颠覆", "改变"),
        ("待验证", ""),
        ("主席修订", ""),
        ("需在提交前补充", ""),
        ("提交前补充", ""),
    )
    for pattern, replacement in replacements:
        if pattern not in cleaned:
            continue
        before = cleaned
        cleaned = cleaned.replace(pattern, replacement)
        contamination_removed.append(
            {
                "category": AUTO_CLEANED_TEXT_CATEGORY,
                "section": section,
                "pattern": pattern,
                "text": before,
            }
        )

    before_preface = cleaned
    cleaned = re.sub(r"^好的，根据[^，。]*(?:[，。])?", "", cleaned).strip()
    if cleaned != before_preface:
        contamination_removed.append(
            {
                "category": AUTO_CLEANED_TEXT_CATEGORY,
                "section": section,
                "pattern": "好的，根据",
                "text": before_preface,
            }
        )

    cleaned = re.sub(r"^[：:，,。；;\s]+", "", cleaned).strip()
    return re.sub(r"\s{2,}", " ", cleaned).strip()


def _contains_yaml_front_matter_evidence_metadata(text: str) -> bool:
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        if lines[index].strip() != "---":
            index += 1
            continue
        index += 1
        while index < len(lines) and lines[index].strip() != "---":
            if YAML_FRONT_MATTER_EVIDENCE_KEY_RE.search(lines[index]):
                return True
            index += 1
        if index < len(lines) and lines[index].strip() == "---":
            index += 1

    first_content_index = next((i for i, line in enumerate(lines) if line.strip()), None)
    if first_content_index is None or not YAML_FRONT_MATTER_EVIDENCE_KEY_RE.search(lines[first_content_index]):
        return False
    for line in lines[first_content_index + 1 :]:
        stripped = line.strip()
        if stripped == "---":
            return True
        if stripped and not line.startswith((" ", "\t")) and not stripped.startswith("-"):
            return False
    return False


def _contains_toml_front_matter_evidence_metadata(text: str) -> bool:
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        if lines[index].strip() != "+++":
            index += 1
            continue
        index += 1
        while index < len(lines) and lines[index].strip() != "+++":
            if TOML_FRONT_MATTER_EVIDENCE_KEY_RE.search(lines[index]):
                return True
            index += 1
        if index < len(lines) and lines[index].strip() == "+++":
            index += 1
    return False


def _contains_ini_section_evidence_metadata(text: str) -> bool:
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        if not INI_EVIDENCE_SECTION_RE.search(lines[index]):
            index += 1
            continue
        index += 1
        while index < len(lines) and not re.match(r"^\s*\[[^\]]+\]\s*$", lines[index]):
            if INI_EVIDENCE_KEY_RE.search(lines[index]):
                return True
            index += 1
    return False


def _contains_csv_evidence_metadata(text: str) -> bool:
    rows = []
    try:
        reader = csv.reader(io.StringIO(text))
        rows = [[cell.strip().strip("\"'") for cell in row] for row in reader]
    except csv.Error:
        return False
    for index, row in enumerate(rows[:-1]):
        if len(row) < 2 or not any(CSV_EVIDENCE_METADATA_HEADER_RE.match(cell) for cell in row):
            continue
        next_row = rows[index + 1]
        if len(next_row) >= 2 and any(cell for cell in next_row):
            return True
    return False


def _contains_fenced_json_evidence_metadata(text: str) -> bool:
    for match in MARKDOWN_CODE_FENCE_RE.finditer(text):
        language = match.group("language").strip().lower()
        body = match.group("body").strip()
        if language != "json" and not body.startswith(("{", "[")):
            continue
        if JSON_EVIDENCE_METADATA_KEY_RE.search(body):
            return True
    return False


def _removal_for_line(line: str, in_fence: bool) -> dict[str, str] | None:
    comparable_line = line.lower()
    if in_fence:
        return {"category": "format_pollution", "pattern": "markdown_fence"}
    if re.search(r"^好的，下面.*撰写", line):
        return {"category": "ai_preface", "pattern": "好的，下面"}
    for pattern in ("support_gap", "support_gaps", "支撑不足提示", "撰写说明"):
        if pattern in comparable_line:
            return {"category": "support_gap", "pattern": pattern}
    if line.startswith("```"):
        return {"category": "format_pollution", "pattern": "markdown_fence"}
    if re.match(r"^#{1,6}\s+", line):
        return {"category": "format_pollution", "pattern": "markdown_heading"}
    if JSON_WRAPPER_RE.match(line):
        return {"category": "format_pollution", "pattern": "json_wrapper"}
    official_field = EMPTY_OFFICIAL_FIELD_RE.match(line)
    if official_field:
        return {"category": "json_wrapper", "pattern": official_field.group(1).lower()}
    official_field_wrapper = OFFICIAL_FIELD_WRAPPER_RE.match(line)
    if official_field_wrapper:
        return {"category": "json_wrapper", "pattern": official_field_wrapper.group(1).lower()}
    internal_field = INTERNAL_FIELD_RE.match(line)
    if internal_field:
        return {"category": "internal_field", "pattern": internal_field.group(1).lower()}
    evidence_metadata_field = EVIDENCE_METADATA_FIELD_RE.match(line)
    if evidence_metadata_field:
        return {"category": "evidence_metadata", "pattern": evidence_metadata_field.group(1).lower()}
    for pattern in ("根据会审策略", "多 Agent 会审", "多Agent会审", "主席汇总", "deliberation", "generation_logs"):
        if pattern.lower() in comparable_line:
            return {"category": "internal_trace", "pattern": pattern}
    if "patents.google.com" in comparable_line:
        return {"category": "evidence_metadata", "pattern": "patents.google.com"}
    for pattern in ("可能不具备创造性", "禁止直接提交", "存在充分公开风险"):
        if pattern in line:
            return {"category": "unfavorable_statement", "pattern": pattern}
    if _looks_like_mermaid(line):
        return {"category": "format_pollution", "pattern": "mermaid"}
    return None


def _looks_like_mermaid(line: str) -> bool:
    mermaid_starters = (
        "flowchart",
        "graph",
        "sequenceDiagram",
        "classDiagram",
        "stateDiagram",
        "erDiagram",
        "gantt",
    )
    if line.startswith(mermaid_starters):
        return True
    return bool(re.search(r"\w+\s*(-->|---|==>|-.->)\s*\w+", line))


def _strip_inline_markdown(line: str) -> str:
    line = re.sub(r"^\s*#{1,6}\s*", "", line)
    return line.strip()


def _parse_figure_plan(drawing_description: str) -> list[OfficialFigurePlanItem]:
    items: list[OfficialFigurePlanItem] = []
    for line in drawing_description.splitlines():
        match = re.match(r"^(图\d+)(?:为|是)(.+?)(?:。|$)", line.strip())
        if not match:
            continue
        figure_no, title = match.groups()
        items.append(
            OfficialFigurePlanItem(
                figure_no=figure_no,
                title=title.strip(),
                description=line.strip(),
                referenced_sections=["drawing_description"],
            )
        )
    return items


def _add_docx_section(doc: Document, heading: str, text: str) -> None:
    doc.add_heading(heading, level=1)
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)


def _dict_item_lines(items: list[dict[str, str]]) -> list[str]:
    if not items:
        return ["- 无"]
    return ["- " + "；".join(f"{key}={value}" for key, value in item.items()) for item in items]


def _utc_now_iso() -> str:
    from datetime import datetime, timezone

    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
