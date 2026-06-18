from __future__ import annotations

import hashlib
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document

from backend.app.schemas import DraftPackage, FilingReadinessIssue, FilingReadinessReport


RULES_VERSION = "filing-readiness-v1"
TECHNICAL_FIELD_SUGGESTION = "建筑信息模型、三维点云处理、计算机视觉与计算机辅助工程量计算技术领域"

_FORMAL_FIELDS = (
    ("abstract", "abstract"),
    ("claims", "claims"),
    ("description", "description"),
    ("drawing_description", "drawings"),
)

_FORMAT_PATTERNS = (
    r"```",
    r"^#{1,6}\s+.+",
    r"\bmermaid\b",
    r"\bflowchart\b",
    r"\bgraph\s+TD\b",
    r"\bsequenceDiagram\b",
    r"^\s*(image_prompt|prompt|diagram)\s*[:：]",
)

_INTERNAL_TRACE_PATTERNS = (
    "多Agent会审",
    "deliberation",
    "generation_logs",
    "根据会审策略",
    "主席汇总失败",
)

_INTERNAL_TRACE_REGEXES = (
    re.compile(r"根据技术交底书\s*(生成|撰写|输出|自动生成)"),
)

_UNFAVORABLE_PATTERNS = (
    "可能不具备创造性",
    "容易被现有技术组合",
    "存在充分公开风险",
    "禁止直接提交",
)

_SUBJECT_MATTER_PATTERNS = (
    "人工智能软件方法领域",
    "智能管理方法",
    "造价规则",
)

_UNVERIFIED_EFFECT_PATTERN = re.compile(r"(提升|降低|提高)\s*\d+(?:\.\d+)?\s*%")

_PRIOR_ART_EFFECT_CONTEXT = (
    "现有技术",
    "对比文件",
    "对比文献",
    "CN",
    "US",
    "EP",
    "JP",
    "WO",
)


def assess_filing_readiness(
    project_id: str,
    package: DraftPackage,
    *,
    verified_effects: bool = False,
    report_id: str | None = None,
) -> FilingReadinessReport:
    """Scan a draft package for content that should not enter a filing draft."""
    issues: list[FilingReadinessIssue] = []

    for field_name, target in _FORMAL_FIELDS:
        text = _value_as_text(getattr(package, field_name, ""))
        issues.extend(_scan_format_pollution(text, target))
        issues.extend(_scan_internal_traces(text, target))
        issues.extend(_scan_unfavorable_statements(text, target))
        issues.extend(_scan_subject_matter_risks(text, target))
        if not verified_effects:
            issues.extend(_scan_unverified_effects(text, target))

    export_text = _export_internal_text(package)
    issues.extend(_scan_internal_field_presence(package))
    issues.extend(_scan_format_pollution(export_text, "export"))
    issues.extend(_scan_internal_traces(export_text, "export"))
    issues.extend(_scan_unfavorable_statements(export_text, "export"))
    issues.extend(_scan_subject_matter_risks(export_text, "export"))
    if not verified_effects:
        issues.extend(_scan_unverified_effects(export_text, "export"))

    status = "clean"
    if any(issue.severity == "high" for issue in issues):
        status = "high_risk"
    elif issues:
        status = "warning"

    package_json = package.model_dump_json()
    package_hash = hashlib.sha256(package_json.encode("utf-8")).hexdigest()

    return FilingReadinessReport(
        id=report_id or uuid.uuid4().hex,
        project_id=project_id,
        draft_package_hash=package_hash,
        status=status,
        rules_version=RULES_VERSION,
        issues=issues,
        created_at=datetime.now(timezone.utc).isoformat(),
    )


def readiness_report_to_markdown(report: FilingReadinessReport) -> str:
    lines = [
        "# FILING_READINESS_REPORT",
        "",
        f"- report_id: {report.id}",
        f"- project_id: {report.project_id}",
        f"- status: {report.status}",
        f"- rules_version: {report.rules_version}",
        f"- draft_package_hash: {report.draft_package_hash}",
        "",
        "## Issues",
    ]

    if not report.issues:
        lines.append("无正式提交阻断项。")
        return "\n".join(lines) + "\n"

    for index, issue in enumerate(report.issues, start=1):
        lines.extend(
            [
                "",
                f"### {index}. {issue.category}",
                f"- severity: {issue.severity}",
                f"- target: {issue.target}",
                f"- matched_text: {issue.matched_text}",
                f"- problem: {issue.message}",
                f"- suggestion: {issue.suggestion}",
                f"- can_auto_clean: {issue.can_auto_clean}",
            ]
        )

    return "\n".join(lines) + "\n"


def official_package_to_markdown(package: DraftPackage) -> str:
    abstract = _clean_official_text(package.abstract)
    claims = _clean_official_text(package.claims)
    description = _clean_official_text(package.description)
    drawing_description = _clean_official_text(package.drawing_description)

    return f"""# {package.title}

## 摘要
{abstract}

## 权利要求书
{claims}

## 说明书
{description}

## 附图说明
{drawing_description}
"""


def export_official_docx(package: DraftPackage, output_path) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(package.title, level=0)
    _add_docx_section(doc, "摘要", _clean_official_text(package.abstract))
    _add_docx_section(doc, "权利要求书", _clean_official_text(package.claims))
    _add_docx_section(doc, "说明书", _clean_official_text(package.description))
    _add_docx_section(doc, "附图说明", _clean_official_text(package.drawing_description))
    doc.save(path)
    return path


def _scan_format_pollution(text: str, target: str) -> list[FilingReadinessIssue]:
    issues = []
    for match in _regex_matches(_FORMAT_PATTERNS, text):
        issues.append(
            FilingReadinessIssue(
                category="format_pollution",
                severity="high",
                target=target,
                matched_text=match,
                message="正式稿包含 Markdown、Mermaid 或绘图提示等格式污染。",
                suggestion="删除格式标记和制图过程内容，仅保留可提交的专利正文。",
                can_auto_clean=True,
            )
        )
    return issues


def _scan_internal_traces(text: str, target: str) -> list[FilingReadinessIssue]:
    literal_matches = list(_literal_matches(_INTERNAL_TRACE_PATTERNS, text))
    regex_matches = [match.group(0) for regex in _INTERNAL_TRACE_REGEXES for match in regex.finditer(text)]
    return [
        FilingReadinessIssue(
            category="internal_trace",
            severity="high",
            target=target,
            matched_text=match,
            message="正式稿包含内部过程痕迹。",
            suggestion="删除过程性表述，仅保留权利要求、说明书和附图说明内容。",
            can_auto_clean=True,
        )
        for match in [*literal_matches, *regex_matches]
    ]


def _scan_unfavorable_statements(text: str, target: str) -> list[FilingReadinessIssue]:
    return [
        FilingReadinessIssue(
            category="unfavorable_statement",
            severity="high",
            target=target,
            matched_text=match,
            message="正式稿包含不利于授权或充分公开的自认表述。",
            suggestion="移除不利评价，改为客观描述技术问题、技术方案和可验证效果。",
            can_auto_clean=False,
        )
        for match in _literal_matches(_UNFAVORABLE_PATTERNS, text)
    ]


def _scan_unverified_effects(text: str, target: str) -> list[FilingReadinessIssue]:
    issues: list[FilingReadinessIssue] = []
    for match in _UNVERIFIED_EFFECT_PATTERN.finditer(text):
        context_start = max(0, match.start() - 36)
        context = text[context_start : match.end() + 18]
        if any(marker in context for marker in _PRIOR_ART_EFFECT_CONTEXT):
            continue
        issues.append(
            FilingReadinessIssue(
                category="unverified_effect",
                severity="medium",
                target=target,
                matched_text=match.group(0),
                message="正式稿包含未经验证的量化效果。",
                suggestion="补充实验或业务验证依据；无法验证时删除具体百分比效果。",
                can_auto_clean=False,
            )
        )
    return issues


def _scan_subject_matter_risks(text: str, target: str) -> list[FilingReadinessIssue]:
    return [
        FilingReadinessIssue(
            category="subject_matter_risk",
            severity="medium",
            target=target,
            matched_text=match,
            message="正式稿的技术领域表述可能被理解为抽象软件或管理规则。",
            suggestion=f"建议调整为：{TECHNICAL_FIELD_SUGGESTION}。",
            can_auto_clean=False,
        )
        for match in _literal_matches(_SUBJECT_MATTER_PATTERNS, text)
    ]


def _scan_internal_field_presence(package: DraftPackage) -> list[FilingReadinessIssue]:
    checks = [
        ("mermaid", package.mermaid, "format_pollution"),
        ("image_prompt", package.image_prompt, "format_pollution"),
        ("generation_logs", package.generation_logs, "internal_trace"),
        ("agent_consensus", package.agent_consensus, "internal_trace"),
        ("strategy_brief", package.strategy_brief, "internal_trace"),
        ("disclosure_summary", package.disclosure_summary, "internal_trace"),
        ("patent_point_summary", package.patent_point_summary, "internal_trace"),
    ]
    return [
        _internal_field_issue(field_name, category)
        for field_name, value, category in checks
        if _has_field_content(value)
    ]


def _internal_field_issue(field_name: str, category: str) -> FilingReadinessIssue:
    if category == "format_pollution":
        return FilingReadinessIssue(
            category="format_pollution",
            severity="high",
            target="export",
            matched_text=field_name,
            message="正式稿导出字段包含制图、提示词或中间格式内容。",
            suggestion="正式提交包中删除该内部字段，仅保留摘要、权利要求书、说明书和附图说明。",
            can_auto_clean=True,
        )
    return FilingReadinessIssue(
        category="internal_trace",
        severity="high",
        target="export",
        matched_text=field_name,
        message="正式稿导出字段包含内部过程痕迹。",
        suggestion="正式提交包中删除该内部字段，仅保留可提交正文。",
        can_auto_clean=True,
    )


def _has_field_content(value: object) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, list):
        return bool(value)
    return True


def _export_internal_text(package: DraftPackage) -> str:
    values: list[str] = [
        package.mermaid,
        package.image_prompt,
        "\n".join(package.generation_logs),
        package.agent_consensus or "",
        package.disclosure_summary or "",
        package.patent_point_summary or "",
    ]
    if package.strategy_brief:
        values.append(package.strategy_brief.model_dump_json(ensure_ascii=False))
    return "\n".join(_value_as_text(value) for value in values)


def _literal_matches(patterns: Iterable[str], text: str) -> list[str]:
    matches: list[str] = []
    for pattern in patterns:
        if pattern in text:
            matches.append(_line_containing(text, pattern))
    return matches


def _line_containing(text: str, pattern: str) -> str:
    for line in text.splitlines():
        if pattern in line:
            return line.strip()
    return pattern


def _clean_official_text(text: str) -> str:
    cleaned_lines: list[str] = []
    in_code_block = False

    for line in text.splitlines():
        if "```" in line:
            in_code_block = not in_code_block
            continue
        if in_code_block or _is_structural_pollution_line(line):
            continue
        cleaned_line = _clean_internal_trace_fragments(line)
        if cleaned_line:
            cleaned_lines.append(cleaned_line)

    return "\n".join(cleaned_lines)


def _is_structural_pollution_line(line: str) -> bool:
    return bool(_regex_matches(_FORMAT_PATTERNS, line))


def _clean_internal_trace_fragments(line: str) -> str:
    cleaned = line
    for pattern in _INTERNAL_TRACE_PATTERNS:
        cleaned = re.sub(
            rf"[，,；;]?\s*{re.escape(pattern)}[^，,；;。]*[，,；;]?",
            "，",
            cleaned,
        )
    return _normalize_punctuation(cleaned)


def _normalize_punctuation(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip()
    cleaned = re.sub(r"，{2,}", "，", cleaned)
    cleaned = re.sub(r"，。", "。", cleaned)
    cleaned = re.sub(r"，$", "。", cleaned)
    return cleaned.strip(" ，,；;")


def _regex_matches(patterns: Iterable[str], text: str) -> list[str]:
    matches: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE | re.MULTILINE):
            matches.append(match.group(0))
    return matches


def _value_as_text(value: object) -> str:
    return "" if value is None else str(value)


def _add_docx_section(doc, heading: str, text: str) -> None:
    doc.add_heading(heading, level=1)
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)
