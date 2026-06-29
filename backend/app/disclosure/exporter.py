from __future__ import annotations

import os
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document

from backend.app.schemas import DisclosurePackage


URL_PATTERN = re.compile(r"https?://\S+")
INTERNAL_METADATA_KEYS = (
    "evidence_id",
    "evidence_refs",
    "research_ledger",
    "generation_logs",
    "provider_diagnostics",
    "revision_ledger",
    "source_ledger",
    "sidecar_notes",
    "internal_only",
    "official_safe_patches",
    "attorney_memo",
    "system_trace",
    "material_id",
    "source_id",
    "source_label",
    "修订记录",
    "检索来源台账",
    "证据编号",
    "材料编号",
    "来源标签",
    "引用来源",
    "引用链接",
    "证据来源",
)
INTERNAL_METADATA_KEY_PATTERN = "|".join(re.escape(key) for key in INTERNAL_METADATA_KEYS)
INTERNAL_METADATA_LINE_RE = re.compile(
    rf"(?:^|\s)(?:[-*+]\s*)?(?:[\"']?)"
    rf"(?:{INTERNAL_METADATA_KEY_PATTERN})"
    rf"(?:[\"']?)\s*[:：=]",
    re.IGNORECASE,
)
INTERNAL_METADATA_JSON_RE = re.compile(
    rf"[\"'](?:{INTERNAL_METADATA_KEY_PATTERN})[\"']\s*[:：=]",
    re.IGNORECASE,
)
MARKDOWN_TABLE_ROW_RE = re.compile(r"^\s*\|(?P<cells>.*)\|\s*$")
MARKDOWN_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")
FENCE_START_RE = re.compile(r"^(?P<indent>\s*)(?P<fence>`{3,}|~{3,})(?P<rest>.*)$")
MARKDOWN_HEADING_RE = re.compile(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$")
INTERNAL_SECTION_HEADINGS = {
    "claim chart",
    "provider diagnostics",
    "diagnostics",
    "mermaid 图",
    "mermaid",
    "绘图提示词",
    "自检结果",
    "生成日志",
    "检索来源台账",
    "引用快照",
    "候选专利点",
    "材料覆盖",
    "前置材料摘要",
    "research ledger",
    "research_ledger",
    "source ledger",
    "source_ledger",
    "revision ledger",
    "revision_ledger",
    "修订记录",
    "generation logs",
    "generation_logs",
    "self check",
    "provider_diagnostics",
    "sidecar",
}


def clean_disclosure_to_markdown(package: DisclosurePackage) -> str:
    body = _clean_export_body_markdown(package.body_markdown)
    appendix = _format_public_prior_art_appendix(package, existing_urls=_extract_normalized_urls(body))
    if not appendix:
        return body
    return f"{body}\n\n{appendix}"


def disclosure_sidecar_to_markdown(package: DisclosurePackage) -> str:
    candidates = "\n".join(
        "\n".join(
            [
                f"- {candidate.id} {candidate.title}：{candidate.innovation}",
                f"  证据状态：{candidate.evidence_status}",
                f"  来源：{candidate.source_type}",
                f"  可行依据：{candidate.feasibility_basis or '未填写'}",
                f"  支撑缺口：{'；'.join(candidate.support_gaps) or '无显式缺口'}",
                f"  护城河评分：{candidate.moat_scores.weighted_total}",
            ]
        )
        for candidate in package.candidates
    )
    claim_charts = "\n".join(
        f"- {candidate.title}｜{chart.prior_art_title}｜差异特征：{'；'.join(chart.differentiating_features) or '暂无'}｜撰写建议：{chart.claim_drafting_advice or '暂无'}"
        for candidate in package.candidates
        for chart in candidate.claim_chart
    )
    prior_art = "\n".join(
        f"- [{hit.source}] {hit.title} {hit.publication_number or ''} {hit.url}\n  摘要：{hit.abstract or '无'}\n  差异：{'；'.join(hit.differentiators) or hit.relevance_summary or '待人工复核'}"
        for hit in package.prior_art_hits
    )
    findings = "\n".join(
        f"- [{finding.severity}] {finding.category}: {finding.message} 建议：{finding.suggestion}"
        for finding in package.self_check_findings
    )
    logs = "\n".join(f"- {log}" for log in package.generation_logs)

    # ---- V1.1: Research source ledger section ----
    ledger_section = _format_ledger_section(package)

    # ---- V1.1: Provider diagnostics section ----
    diagnostics_section = _format_diagnostics_section(package)

    # ---- V1.1: Research confidence badge ----
    confidence_badges = {"low": "🔴 低", "medium": "🟡 中", "high": "🟢 高"}
    confidence_badge = confidence_badges.get(package.research_confidence, "⚪ 未知")
    confidence_note = (
        "\n> **检索置信度**：{badge}\n>\n"
        "> 低置信度表示未检索到可引用的公开现有技术文献；交底书不隐含高专利性判断。\n"
        .format(badge=confidence_badge)
    )

    return f"""# {package.title}

## 前置材料摘要
{package.summary}

{confidence_note}
## 材料覆盖
{package.materials_summary}

## 候选专利点
{candidates or "暂无。"}

## Claim Chart
{claim_charts or "暂无。"}

## 公开现有技术
{prior_art or "暂无可用公开检索结果。"}

## 现有技术差异
{package.prior_art_differences}
{ledger_section}
{diagnostics_section}
## 技术交底书
{package.body_markdown}

## Mermaid 图
```mermaid
{package.mermaid}
```

## 绘图提示词
{package.image_prompt}

## 自检结果
{findings or "暂无。"}

## 生成日志
{logs or "暂无。"}
"""


def disclosure_to_markdown(package: DisclosurePackage) -> str:
    return disclosure_sidecar_to_markdown(package)


def export_disclosure_docx(package: DisclosurePackage, output_path: Path, run_dir: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    run_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(package.title, level=0)
    clean_body = _clean_export_body_markdown(package.body_markdown)
    for line in clean_body.splitlines() or [""]:
        doc.add_paragraph(line)
    appendix = _format_public_prior_art_appendix(package, existing_urls=_extract_normalized_urls(clean_body))
    if appendix:
        for line in appendix.splitlines():
            doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def write_disclosure_artifacts(package: DisclosurePackage, run_dir: Path) -> dict[str, Path]:
    run_dir.mkdir(parents=True, exist_ok=True)
    md_path = run_dir / "disclosure.md"
    sidecar_path = run_dir / "disclosure-sidecar.md"
    mmd_path = run_dir / "diagram.mmd"
    prompt_path = run_dir / "image-prompt.md"
    docx_path = run_dir / "disclosure.docx"
    md_path.write_text(clean_disclosure_to_markdown(package), encoding="utf-8")
    sidecar_path.write_text(disclosure_sidecar_to_markdown(package), encoding="utf-8")
    mmd_path.write_text(package.mermaid, encoding="utf-8")
    prompt_path.write_text(package.image_prompt, encoding="utf-8")
    export_disclosure_docx(package, docx_path, run_dir)
    return {"md": md_path, "sidecar": sidecar_path, "mmd": mmd_path, "prompt": prompt_path, "docx": docx_path}


def _render_mermaid(mermaid: str, run_dir: Path) -> Path | None:
    mmdc = shutil.which("mmdc")
    command: list[str] | None = None
    if mmdc:
        command = [mmdc]
    elif os.environ.get("PATENTS_AGENT_ENABLE_NPX_MERMAID") == "1" and shutil.which("npx"):
        command = ["npx", "-y", "@mermaid-js/mermaid-cli", "mmdc"]
    if not command:
        return None
    source = run_dir / "diagram-source.mmd"
    output = run_dir / "diagram.png"
    source.write_text(mermaid, encoding="utf-8")
    try:
        subprocess.run(
            [*command, "-i", str(source), "-o", str(output), "-b", "white"],
            check=True,
            capture_output=True,
            timeout=12,
        )
    except Exception:
        return None
    return output if output.exists() else None


def _add_section(doc: Document, heading: str, text: str) -> None:
    doc.add_heading(heading, level=1)
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)


def _clean_export_body_markdown(body_markdown: str) -> str:
    raw_lines = body_markdown.strip().splitlines()
    if not raw_lines:
        return ""

    raw_lines = _drop_internal_front_matter(raw_lines)
    lines: list[str] = []
    skipping_section_level: int | None = None
    index = 0
    while index < len(raw_lines):
        line = raw_lines[index]
        heading = MARKDOWN_HEADING_RE.match(line)
        if heading:
            level = len(heading.group(1))
            title = _normalize_internal_heading(heading.group(2))
            if skipping_section_level is not None and level <= skipping_section_level:
                skipping_section_level = None
            if title in INTERNAL_SECTION_HEADINGS:
                skipping_section_level = level
                index += 1
                continue
        if skipping_section_level is not None:
            index += 1
            continue
        fence_start = FENCE_START_RE.match(line)
        if fence_start:
            block, next_index = _collect_fenced_block(raw_lines, index, fence_start.group("fence"))
            if not _block_contains_internal_metadata(block):
                lines.extend(entry.rstrip() for entry in block)
            index = next_index
            continue
        if MARKDOWN_TABLE_ROW_RE.match(line):
            block, next_index = _collect_table_block(raw_lines, index)
            lines.extend(_clean_markdown_table_block(block))
            index = next_index
            continue
        if _line_contains_internal_metadata(line):
            index += 1
            continue
        lines.append(line.rstrip())
        index += 1
    return "\n".join(lines).strip()


def _normalize_internal_heading(heading: str) -> str:
    normalized = re.sub(r"\s+", " ", heading.strip()).strip("：:").casefold()
    return normalized


def _drop_internal_front_matter(lines: list[str]) -> list[str]:
    if not lines or lines[0].strip() != "---":
        return lines

    closing_index: int | None = None
    for index in range(1, len(lines)):
        if lines[index].strip() in {"---", "..."}:
            closing_index = index
            break
    if closing_index is None:
        return lines

    block = lines[: closing_index + 1]
    if any(_line_contains_internal_metadata(line) for line in block[1:-1]):
        return lines[closing_index + 1 :]
    return lines


def _collect_fenced_block(lines: list[str], start_index: int, opening_fence: str) -> tuple[list[str], int]:
    block = [lines[start_index]]
    index = start_index + 1
    fence_char = opening_fence[0]
    minimum_width = len(opening_fence)
    closing_re = re.compile(rf"^\s*{re.escape(fence_char)}{{{minimum_width},}}\s*$")
    while index < len(lines):
        block.append(lines[index])
        if closing_re.match(lines[index]):
            return block, index + 1
        index += 1
    return block, index


def _collect_table_block(lines: list[str], start_index: int) -> tuple[list[str], int]:
    block: list[str] = []
    index = start_index
    while index < len(lines) and MARKDOWN_TABLE_ROW_RE.match(lines[index]):
        block.append(lines[index])
        index += 1
    return block, index


def _clean_markdown_table_block(block: list[str]) -> list[str]:
    cleaned: list[str] = []
    data_row_count = 0
    for line in block:
        if MARKDOWN_TABLE_SEPARATOR_RE.match(line):
            cleaned.append(line.rstrip())
            continue
        if _table_row_contains_internal_metadata(line):
            continue
        cleaned.append(line.rstrip())
        if not _looks_like_table_header(line):
            data_row_count += 1
    if data_row_count == 0:
        return []
    return cleaned


def _looks_like_table_header(line: str) -> bool:
    if MARKDOWN_TABLE_SEPARATOR_RE.match(line):
        return False
    match = MARKDOWN_TABLE_ROW_RE.match(line)
    if not match:
        return False
    cells = [cell.strip() for cell in match.group("cells").split("|")]
    normalized_cells = {_normalize_internal_heading(cell) for cell in cells if cell.strip()}
    return normalized_cells <= {"字段", "内容", "field", "value", "label", "key"} or normalized_cells == {"---"}


def _block_contains_internal_metadata(lines: list[str]) -> bool:
    return any(_line_contains_internal_metadata(line) or _table_row_contains_internal_metadata(line) for line in lines)


def _line_contains_internal_metadata(line: str) -> bool:
    return bool(INTERNAL_METADATA_LINE_RE.search(line) or INTERNAL_METADATA_JSON_RE.search(line))


def _table_row_contains_internal_metadata(line: str) -> bool:
    match = MARKDOWN_TABLE_ROW_RE.match(line)
    if not match:
        return False
    cells = [cell.strip() for cell in match.group("cells").split("|")]
    if not cells:
        return False
    key = _normalize_internal_heading(cells[0])
    return any(key == candidate.casefold() for candidate in INTERNAL_METADATA_KEYS)


def _format_public_prior_art_appendix(package: DisclosurePackage, *, existing_urls: set[str] | None = None) -> str:
    if not package.prior_art_hits:
        return ""

    existing_urls = existing_urls or set()
    lines = ["## 公开现有技术链接", ""]
    for hit in package.prior_art_hits:
        url = _normalize_url(hit.url or "")
        if not url or url in existing_urls:
            continue
        label = hit.title
        if hit.publication_number:
            label = f"{label}（{hit.publication_number}）"
        lines.append(f"- {label}: {url}")
    return "\n".join(lines) if len(lines) > 2 else ""


def _extract_normalized_urls(text: str) -> set[str]:
    return {_normalize_url(match.group(0)) for match in URL_PATTERN.finditer(text) if _normalize_url(match.group(0))}


def _normalize_url(url: str) -> str:
    return url.strip().rstrip(".,;:)]}>。！？）】")


def _format_ledger_section(package: "DisclosurePackage") -> str:
    """Format the research source ledger as a markdown section."""
    ledger = package.research_ledger
    if not ledger:
        return ""

    lines = [
        "## 检索来源台账",
        "",
    ]
    entries = ledger.get("entries", [])
    if not entries:
        lines.append("暂无检索记录。")
        lines.append("")
        return "\n".join(lines)

    lines.extend([
        f"- 总命中数：{ledger.get('total_hits', 0)}",
        f"- 总引用数：{ledger.get('total_citations', 0)}",
        "",
        "| 来源 | 类型 | 检索词 | 状态 | 命中 | 保留 | 失败原因 |",
        "|------|------|--------|------|------|------|----------|",
    ])
    for entry in entries:
        provider = entry.get("provider", "")
        kind = entry.get("kind", "")
        query = (entry.get("query", "") or "")[:50]
        status = entry.get("status", "running")
        hit_count = entry.get("hit_count", 0)
        retained = entry.get("retained_count", 0)
        failure = (entry.get("failure_reason", "") or "")[:60]

        status_icon = {"ok": "✅", "failed": "❌", "timeout": "⏰", "skipped": "⏭️", "running": "🔄"}.get(status, "❓")
        lines.append(
            f"| {provider} | {kind} | {query} | {status_icon} {status} | {hit_count} | {retained} | {failure} |"
        )

    # Citation snapshots
    all_citations: list[dict] = []
    for entry in entries:
        citations = entry.get("citations", [])
        if isinstance(citations, list):
            all_citations.extend(citations)

    if all_citations:
        lines.append("")
        lines.append("### 引用快照")
        lines.append("")
        for i, cit in enumerate(all_citations[:20], start=1):
            pub = cit.get("publication_number", "") or ""
            title = cit.get("title", "")[:80] or ""
            source = cit.get("source", "") or ""
            abstract = cit.get("abstract_snippet", "")[:60] or ""
            lines.append(f"{i}. [{source}] {pub} {title}")
            if abstract:
                lines.append(f"   {abstract}")

    lines.append("")
    return "\n".join(lines)


def _format_diagnostics_section(package: "DisclosurePackage") -> str:
    """Format provider diagnostics as a markdown section."""
    diagnostics = package.provider_diagnostics
    if not diagnostics:
        return ""

    lines = [
        "## 检索链路诊断",
        "",
    ]
    for diag in diagnostics:
        phase = diag.get("phase", "unknown")
        phase_label = "🔍 检索前" if phase == "pre_flight" else "📊 检索后"
        lines.append(f"### {phase_label}")
        lines.append("")

        available = diag.get("available_providers", [])
        if available:
            lines.append(f"- 可用来源：{'、'.join(available)}")
        else:
            lines.append("- 可用来源：无")

        skipped = diag.get("skipped_providers", [])
        if skipped:
            lines.append("- 跳过来源：")
            for skip in skipped:
                provider = skip.get("provider", "") if isinstance(skip, dict) else str(skip)
                reason = skip.get("reason", "") if isinstance(skip, dict) else ""
                lines.append(f"  - {provider}：{reason}")

        warnings_list = diag.get("warnings", [])
        if warnings_list:
            lines.append("- 警告：")
            for w in warnings_list:
                lines.append(f"  - {w}")

        lines.append("")
    return "\n".join(lines)
