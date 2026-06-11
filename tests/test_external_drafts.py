import pytest
from docx import Document
from pydantic import ValidationError

import backend.app.external_drafts as external_drafts
from backend.app.schemas import (
    DraftPackage,
    ExternalDraftIntakeRun,
    ExternalDraftReviewBundle,
    ExternalDraftSource,
    ExternalDraftSourceCreate,
    IntakeIssue,
    SectionConfidence,
    SectionConfidenceItem,
)


def test_external_draft_models_capture_source_intake_and_confidence():
    source = ExternalDraftSource(
        id="src-1",
        project_id="project-1",
        source_type="pasted_text",
        file_name="pasted.txt",
        content_hash="hash-raw",
        raw_text="权利要求书\n1. 一种方法。\n说明书\n本发明涉及数据处理。",
        raw_path="",
        metadata={"input": "paste"},
    )
    confidence = SectionConfidence(
        title=SectionConfidenceItem(score=0.4, source_markers=[], warnings=["未识别发明名称"]),
        abstract=SectionConfidenceItem(score=0.0, source_markers=[], warnings=["未识别摘要"]),
        claims=SectionConfidenceItem(score=0.95, source_markers=["权利要求书"], warnings=[]),
        description=SectionConfidenceItem(score=0.9, source_markers=["说明书"], warnings=[]),
        drawing_description=SectionConfidenceItem(score=0.0, source_markers=[], warnings=["未识别附图说明"]),
    )
    issue = IntakeIssue(
        id="intake-1",
        category="missing_section",
        severity="medium",
        section="abstract",
        message="未识别摘要章节。",
        suggested_action="在章节确认界面补充摘要。",
        blocks_quality_run=False,
    )
    run = ExternalDraftIntakeRun(
        id="run-1",
        project_id="project-1",
        source_id=source.id,
        status="needs_review",
        parser_version="external-draft-parser-v1",
        source_hash=source.content_hash,
        parsed_package=DraftPackage(
            title="未命名发明",
            abstract="",
            claims="1. 一种方法。",
            description="本发明涉及数据处理。",
            drawing_description="",
            mermaid="",
            image_prompt="",
            review_findings=[],
            citations=[],
            generation_logs=[],
        ),
        section_confidence=confidence,
        intake_issues=[issue],
        unassigned_fragments=[],
        working_draft_hash="hash-working",
    )

    create = ExternalDraftSourceCreate(source_type="pasted_text", text=source.raw_text, file_name="pasted.txt")

    assert create.source_type == "pasted_text"
    assert source.content_hash == "hash-raw"
    assert run.status == "needs_review"
    assert run.section_confidence.claims.score == 0.95
    assert run.intake_issues[0].category == "missing_section"


@pytest.mark.parametrize(
    "payload",
    [
        {"source_type": "pasted_text", "text": ""},
        {"source_type": "pasted_text", "text": "   "},
        {"source_type": "markdown_file", "text": "", "file_content": ""},
        {"source_type": "markdown_file", "text": "  ", "file_content": "\n\t"},
        {"source_type": "docx_file", "text": "", "file_content": ""},
        {"source_type": "docx_file", "text": "  ", "file_content": "\n\t"},
    ],
)
def test_external_draft_source_create_rejects_empty_content_for_source_type(payload):
    with pytest.raises(ValidationError):
        ExternalDraftSourceCreate(**payload)


def test_external_draft_review_bundle_scores_are_bounded_when_present():
    assert ExternalDraftReviewBundle(project_id="project-1", initial_score=0, latest_score=100).latest_score == 100

    with pytest.raises(ValidationError):
        ExternalDraftReviewBundle(project_id="project-1", initial_score=-1)

    with pytest.raises(ValidationError):
        ExternalDraftReviewBundle(project_id="project-1", latest_score=999)


from backend.app.external_drafts import (
    create_external_draft_source,
    extract_docx_text,
    external_draft_review_bundle_to_markdown,
    parse_external_draft_source,
    working_draft_hash,
)


def test_docx_external_draft_text_extraction(tmp_path):
    docx_path = tmp_path / "external-draft.docx"
    document = Document()
    document.add_heading("一种DOCX外部稿处理方法", level=1)
    document.add_paragraph("摘要")
    document.add_paragraph("本发明公开一种DOCX外部稿处理方法。")
    document.add_paragraph("权利要求书")
    document.add_paragraph("1. 一种方法，其特征在于，读取DOCX段落并生成工作稿。")
    document.add_paragraph("说明书")
    document.add_paragraph("本发明涉及文档解析。")
    document.save(docx_path)

    text = extract_docx_text(docx_path)

    assert "一种DOCX外部稿处理方法" in text
    assert "权利要求书" in text
    assert "读取DOCX段落" in text


def test_markdown_external_draft_parses_into_working_package():
    source = create_external_draft_source(
        project_id="project-1",
        source_type="markdown_file",
        text=(
            "# 一种指标缺口驱动的无人机采集方法\n\n"
            "## 摘要\n"
            "本发明公开一种按指标缺口生成无人机采集任务的方法。\n\n"
            "## 权利要求书\n"
            "1. 一种无人机采集方法，其特征在于，计算指标证据缺失度并生成采集任务包。\n"
            "2. 根据权利要求1所述的方法，其特征在于，按置信度增益排序任务。\n\n"
            "## 说明书\n"
            "技术领域\n"
            "本发明涉及城市体检数据采集。\n"
            "具体实施方式\n"
            "系统计算指标证据缺失度、传感器可达性和采集窗口。\n\n"
            "## 附图说明\n"
            "图1为方法流程图。\n"
        ),
        file_name="draft.md",
    )

    run = parse_external_draft_source(project_id="project-1", source=source)

    assert run.status == "completed"
    assert run.parsed_package is not None
    assert run.parsed_package.title == "一种指标缺口驱动的无人机采集方法"
    assert "指标证据缺失度" in run.parsed_package.claims
    assert run.section_confidence is not None
    assert run.section_confidence.claims.score >= 0.9
    assert run.working_draft_hash == working_draft_hash(run.parsed_package)


def test_markdown_h1_title_heading_uses_following_line_as_title():
    source = create_external_draft_source(
        project_id="project-1",
        source_type="markdown_file",
        text=(
            "# 发明名称\n"
            "一种数据处理方法\n\n"
            "## 权利要求书\n"
            "1. 一种数据处理方法，其特征在于，执行数据清洗和结果输出。\n\n"
            "## 说明书\n"
            "本发明涉及数据处理。\n"
        ),
        file_name="draft.md",
    )

    run = parse_external_draft_source(project_id="project-1", source=source)

    assert run.parsed_package is not None
    assert run.parsed_package.title == "一种数据处理方法"
    assert "一种数据处理方法" not in run.unassigned_fragments


def test_bare_name_inside_description_does_not_reset_title_section():
    source = create_external_draft_source(
        project_id="project-1",
        source_type="pasted_text",
        text=(
            "发明名称\n"
            "一种数据处理方法\n"
            "权利要求书\n"
            "1. 一种数据处理方法，其特征在于，处理字段名称。\n"
            "说明书\n"
            "本发明涉及数据处理。\n"
            "名称\n"
            "该字段名称用于标识数据列。\n"
        ),
        file_name="pasted.txt",
    )

    run = parse_external_draft_source(project_id="project-1", source=source)

    assert run.parsed_package is not None
    assert run.parsed_package.title == "一种数据处理方法"
    assert "名称\n该字段名称用于标识数据列。" in run.parsed_package.description


def test_spaced_chinese_headings_are_recognized():
    source = create_external_draft_source(
        project_id="project-1",
        source_type="pasted_text",
        text=(
            "发明名称\n"
            "一种数据处理方法\n"
            "【摘 要】\n"
            "本发明公开一种数据处理方法。\n"
            "**权 利 要 求 书：**\n"
            "1. 一种数据处理方法，其特征在于，执行数据清洗。\n"
            "（说 明 书）\n"
            "本发明涉及数据处理。\n"
        ),
        file_name="pasted.txt",
    )

    run = parse_external_draft_source(project_id="project-1", source=source)

    assert run.status == "completed"
    assert run.parsed_package is not None
    assert run.parsed_package.abstract == "本发明公开一种数据处理方法。"
    assert "执行数据清洗" in run.parsed_package.claims
    assert run.parsed_package.description == "本发明涉及数据处理。"


def test_inline_heading_content_is_parsed_into_sections():
    source = create_external_draft_source(
        project_id="project-1",
        source_type="pasted_text",
        text=(
            "发明名称：一种数据处理方法\n"
            "摘要：本发明公开一种数据处理方法。\n"
            "权利要求书：1. 一种数据处理方法，其特征在于，执行数据清洗。\n"
            "说明书：本发明涉及数据处理。\n"
            "附图说明：图1为方法流程图。\n"
        ),
        file_name="pasted.txt",
    )

    run = parse_external_draft_source(project_id="project-1", source=source)

    assert run.status == "completed"
    assert run.parsed_package is not None
    assert run.parsed_package.title == "一种数据处理方法"
    assert run.parsed_package.abstract == "本发明公开一种数据处理方法。"
    assert run.parsed_package.claims == "1. 一种数据处理方法，其特征在于，执行数据清洗。"
    assert run.parsed_package.description == "本发明涉及数据处理。"
    assert run.parsed_package.drawing_description == "图1为方法流程图。"
    assert run.unassigned_fragments == []


def test_external_draft_parser_programmer_errors_surface(monkeypatch):
    source = create_external_draft_source(
        project_id="project-1",
        source_type="pasted_text",
        text="发明名称\n一种数据处理方法\n权利要求书\n1. 一种方法。\n说明书\n本发明涉及数据处理。",
        file_name="pasted.txt",
    )

    def raise_runtime_error(raw_text):
        raise RuntimeError(f"boom: {raw_text[:4]}")

    monkeypatch.setattr(external_drafts, "parse_sections", raise_runtime_error)

    with pytest.raises(RuntimeError, match="boom"):
        parse_external_draft_source(project_id="project-1", source=source)


def test_external_draft_needs_review_when_claims_are_missing():
    source = create_external_draft_source(
        project_id="project-1",
        source_type="pasted_text",
        text="发明名称\n一种数据处理方法\n说明书\n本发明涉及数据处理。",
        file_name="pasted.txt",
    )

    run = parse_external_draft_source(project_id="project-1", source=source)

    assert run.status == "needs_review"
    assert any(issue.category == "missing_section" and issue.section == "claims" for issue in run.intake_issues)
    assert any(issue.blocks_quality_run for issue in run.intake_issues)


def test_external_draft_flags_duplicate_sections_and_malformed_claim_numbering():
    source = create_external_draft_source(
        project_id="project-1",
        source_type="pasted_text",
        text=(
            "发明名称\n一种处理方法\n"
            "摘要\n摘要一。\n"
            "摘要\n摘要二。\n"
            "权利要求书\n一、一种处理方法。\n"
            "说明书\n本发明涉及数据处理。\n"
        ),
        file_name="pasted.txt",
    )

    run = parse_external_draft_source(project_id="project-1", source=source)

    assert any(issue.category == "duplicate_section" and issue.section == "abstract" for issue in run.intake_issues)
    assert any(issue.category == "malformed_claim_numbering" for issue in run.intake_issues)
