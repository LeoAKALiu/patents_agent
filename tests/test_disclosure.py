from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.app.disclosure.exporter import export_disclosure_docx
from backend.app.disclosure.generator import DisclosureGenerator
from backend.app.disclosure.material_parser import read_project_material_text
from backend.app.disclosure.prior_art import StaticPriorArtProvider, parse_cnipa_epub_html
from backend.app.llm import FakeLLMClient
from backend.app.main import create_app
from backend.app.patent_mode import UTILITY_MODEL_MODE_PREFIX
from backend.app.schemas import (
    ClaimChartItem,
    DeliberationRun,
    DeliberationStageResult,
    DisclosurePackage,
    PatentPointCandidate,
    PatentStrategyBrief,
    PriorArtHit,
    ProjectMaterial,
    ProjectRecord,
)


def test_project_material_parser_reads_text_docx_pptx_and_rejects_blank_pdf(tmp_path: Path):
    text_path = tmp_path / "draft.txt"
    text_path.write_text("一种基于神经网络的图像缺陷识别方法，包含采集、训练和输出。", encoding="utf-8")
    text, warnings = read_project_material_text(text_path)
    assert "图像缺陷识别" in text
    assert warnings == []

    docx_path = tmp_path / "design.docx"
    doc = Document()
    doc.add_paragraph("系统包括采集模块、模型训练模块和缺陷输出模块。")
    doc.save(docx_path)
    docx_text, _ = read_project_material_text(docx_path)
    assert "模型训练模块" in docx_text

    pptx_path = tmp_path / "deck.pptx"
    with zipfile.ZipFile(pptx_path, "w") as archive:
        archive.writestr(
            "ppt/slides/slide1.xml",
            '<p:sld xmlns:p="p" xmlns:a="a"><a:t>流程包括图像采集</a:t><a:t>缺陷定位输出</a:t></p:sld>',
        )
    pptx_text, _ = read_project_material_text(pptx_path)
    assert "缺陷定位输出" in pptx_text

    blank_pdf = tmp_path / "blank.pdf"
    fitz = pytest.importorskip("fitz")
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(blank_pdf)
    pdf.close()
    with pytest.raises(ValueError, match="no text layer"):
        read_project_material_text(blank_pdf)


def test_project_material_parser_rejects_empty_text_and_invalid_docx(tmp_path: Path):
    empty_text = tmp_path / "empty.md"
    empty_text.write_text(" \n\t", encoding="utf-8")
    with pytest.raises(ValueError, match="empty|文件为空|文本为空|no extractable text"):
        read_project_material_text(empty_text)

    fake_docx = tmp_path / "fake.docx"
    fake_docx.write_bytes(b"not a real docx archive")
    with pytest.raises(ValueError, match="DOCX|docx|格式"):
        read_project_material_text(fake_docx)


def test_parse_cnipa_epub_html_extracts_abstract_and_dedupes():
    html = """
    <div class="item">
      <h1 class="title">一种图像缺陷检测方法</h1>
      <div class="qrcode" title="http://epub.cnipa.gov.cn/patent/CN123456789A"></div>
      <dt>申请公布号：</dt><dd>CN123456789A</dd>
      <dt>摘要：</dt><dd>本发明涉及图像缺陷检测。</dd>
    </div>
    <div class="item">
      <h1 class="title">一种图像缺陷检测方法</h1>
      <div class="qrcode" title="http://epub.cnipa.gov.cn/patent/CN123456789A"></div>
      <dt>申请公布号：</dt><dd>CN123456789A</dd>
    </div>
    """
    hits = parse_cnipa_epub_html(html, "图像 缺陷")
    assert len(hits) == 1
    assert hits[0].publication_number == "CN123456789A"
    assert hits[0].abstract == "本发明涉及图像缺陷检测。"


def test_disclosure_research_prompts_include_structured_project_metadata():
    llm = FakeLLMClient(
        {
            "disclosure_scan": '{"summary":"多模态道路病害检测","materials_summary":"无","technical_keywords":["道路病害"],"implementation_gaps":[]}',
            "patent_points": '{"candidates":[{"id":"p1","title":"跨模态病害检测","technical_problem":"夜间检测误检率高","innovation":"跨模态特征对齐","technical_solution":"融合可见光、红外和激光雷达","beneficial_effects":["夜间准确率提升"],"protection_focus":["方法"],"grantability_score":0.8,"rationale":"完整"}],"selected_candidate_id":"p1"}',
            "prior_art_terms": '["道路病害 多模态 检测"]',
            "disclosure_body": "# 交底书",
            "disclosure_mermaid": "flowchart TD",
            "disclosure_image_prompt": "黑白线稿",
            "disclosure_self_check": "[]",
        }
    )
    generator = DisclosureGenerator(llm, StaticPriorArtProvider())
    project = ProjectRecord(
        id="project-meta",
        name="道路病害检测",
        draft_text="一种道路病害检测方法。",
        technical_field="计算机视觉、市政工程检测",
        background="人工巡检效率低且夜间漏检。",
        pain_point="光照变化导致误检率高。",
        technical_solution="可见光、红外和激光雷达多模态融合。",
        innovation="跨模态特征对齐。",
        embodiments="车载多传感器在经十路巡检。",
        beneficial_effects="夜间检测准确率提升。",
    )

    generator.generate(
        project=project,
        materials=[],
        context_chunks=[],
        max_prior_art_results=0,
    )

    prompts = {call.stage: call.user_prompt for call in llm.calls}
    assert "计算机视觉、市政工程检测" in prompts["disclosure_scan"]
    assert "跨模态特征对齐" in prompts["prior_art_terms"]
    assert "夜间检测准确率提升" in prompts["prior_art_terms"]


def test_disclosure_generator_runs_pipeline_and_records_prior_art():
    llm = _disclosure_llm()
    provider = StaticPriorArtProvider(hits=[_prior_art_hit()])
    generator = DisclosureGenerator(llm, provider)
    project = ProjectRecord(id="p1", name="图像缺陷识别", draft_text="一种基于神经网络的图像缺陷识别方法。")

    package, stage_results, warnings = generator.generate(
        project=project,
        materials=[],
        context_chunks=[],
        max_prior_art_results=8,
        user_candidates=[],
    )

    assert [call.stage for call in llm.calls] == [
        "disclosure_scan",
        "patent_points",
        "prior_art_terms",
        "prior_art_relevance",
        "disclosure_body",
        "disclosure_mermaid",
        "disclosure_image_prompt",
        "disclosure_self_check",
    ]
    assert package.selected_candidate_id == "p1"
    assert package.prior_art_hits[0].url == "https://patents.google.com/patent/CN123456789A"
    assert package.research_ledger["entries"][0]["provider"] == "static_prior_art"
    assert package.research_confidence == "medium"
    assert stage_results[0]["phase"] == "project_scan"
    assert warnings == []


def test_disclosure_generator_dedupes_user_and_generated_candidate_ids():
    llm = _disclosure_llm(
        extra={
            "patent_points": '{"candidates":[{"id":"user-point","title":"模型重复点","technical_problem":"模型问题","innovation":"模型创新","technical_solution":"模型方案"}],"selected_candidate_id":"user-point"}',
        }
    )
    generator = DisclosureGenerator(llm, StaticPriorArtProvider())
    project = ProjectRecord(id="p1", name="外立面逆建模", draft_text="一种既有建筑外立面逆建模方法。")
    first_user = _user_candidate("user-point", "用户第一专利点")
    duplicate_user = _user_candidate("user-point", "用户重复专利点")

    package, _, _ = generator.generate(
        project=project,
        materials=[],
        context_chunks=[],
        max_prior_art_results=0,
        user_candidates=[first_user, duplicate_user],
    )

    assert [candidate.id for candidate in package.candidates] == ["user-point"]
    assert [candidate.title for candidate in package.candidates] == ["用户第一专利点"]
    assert package.selected_candidate_id == "user-point"


def test_disclosure_generator_injects_user_candidate_context_into_prompts():
    llm = _disclosure_llm()
    generator = DisclosureGenerator(llm, StaticPriorArtProvider())
    project = ProjectRecord(id="p1", name="外立面逆建模", draft_text="一种既有建筑外立面逆建模方法。")
    user_candidate = _user_candidate("user-point", "遮挡洞口语义补全", selected=True)

    generator.generate(
        project=project,
        materials=[],
        context_chunks=[],
        max_prior_art_results=0,
        user_candidates=[user_candidate],
    )

    prompts_by_stage = {call.stage: call.user_prompt for call in llm.calls}
    assert "遮挡洞口语义补全" in prompts_by_stage["patent_points"]
    assert "遮挡洞口语义补全" in prompts_by_stage["prior_art_terms"]


def test_disclosure_generator_utility_model_prompts_prioritize_structure():
    llm = _disclosure_llm(
        extra={
            "patent_points": '{"candidates":[{"id":"p1","title":"可调安装支架结构","technical_problem":"支架角度调节不便","innovation":"限位件与支撑臂配合","technical_solution":"底座、支撑臂和限位件形成可调连接"}],"selected_candidate_id":"p1"}',
        }
    )
    generator = DisclosureGenerator(llm, StaticPriorArtProvider())
    project = ProjectRecord(
        id="p1",
        name="可调安装支架",
        draft_text=f"{UTILITY_MODEL_MODE_PREFIX}专利类型：实用新型。\n一种可调安装支架，包括底座、支撑臂和限位件。",
    )

    package, _, _ = generator.generate(
        project=project,
        materials=[],
        context_chunks=[],
        max_prior_art_results=0,
        user_candidates=[],
    )

    prompts_by_stage = {call.stage: call.user_prompt for call in llm.calls}
    assert "中国实用新型专利" in llm.calls[0].system_prompt
    assert "候选结构点" in prompts_by_stage["patent_points"]
    assert "各部件连接/安装/配合关系" in prompts_by_stage["disclosure_body"]
    assert package.candidates[0].protection_focus == ["结构", "部件", "连接关系"]


def test_disclosure_api_lifecycle_and_generation_injection(tmp_path: Path):
    llm = _disclosure_llm(
        extra={
            "claims": "1. 一种图像缺陷识别方法，其特征在于，包括采集图像并输出缺陷位置。",
            "description": "技术领域\n本发明涉及AI检测。\n发明内容\n本发明基于交底书限定输入输出。",
            "abstract": "本发明公开了一种图像缺陷识别方法。",
            "drawings": "图1为方法流程图。",
            "diagram": "flowchart TD\nA[采集] --> B[输出]",
            "image_prompt": "黑白线稿，展示采集和输出流程。",
        }
    )
    client = TestClient(
        create_app(
            data_dir=tmp_path,
            llm_client=llm,
            prior_art_provider=StaticPriorArtProvider(hits=[_prior_art_hit()]),
        )
    )
    project_response = client.post(
        "/api/projects",
        json={"name": "图像缺陷识别", "draft_text": "一种基于神经网络的图像缺陷识别方法。"},
    )
    project_id = project_response.json()["id"]

    material_response = client.post(
        f"/api/projects/{project_id}/materials",
        files={"file": ("material.txt", "补充材料：系统包含采集模块和缺陷定位模块。".encode("utf-8"), "text/plain")},
    )
    assert material_response.status_code == 200
    assert material_response.json()["status"] == "processed"

    run_response = client.post(f"/api/projects/{project_id}/disclosures", json={"trace": False, "max_prior_art_results": 8})
    assert run_response.status_code == 200
    run = run_response.json()
    assert run["status"] == "completed"
    assert run["package"]["prior_art_hits"][0]["publication_number"] == "CN123456789A"

    list_response = client.get(f"/api/projects/{project_id}/disclosures")
    assert list_response.json()["runs"][0]["id"] == run["id"]
    _create_completed_deliberation(client, project_id)

    generate_response = client.post(f"/api/projects/{project_id}/generate", json={})
    assert generate_response.status_code == 200
    package = generate_response.json()
    assert package["disclosure_run_id"] == run["id"]
    assert package["patent_point_summary"] == "图像缺陷识别方法及系统"
    assert any("disclosure: injected" in log for log in package["generation_logs"])

    export_response = client.get(f"/api/projects/{project_id}/disclosures/{run['id']}/export.md")
    assert export_response.status_code == 200
    assert "公开现有技术" in export_response.text
    assert "检索来源台账" in export_response.text
    assert "检索前" in export_response.text
    assert "检索后" in export_response.text
    assert "候选专利点" in export_response.text
    assert "证据状态" in export_response.text


def test_project_material_upload_rejects_empty_text_and_invalid_docx_without_persisting(tmp_path: Path):
    client = TestClient(create_app(data_dir=tmp_path, llm_client=FakeLLMClient({}), load_env_file=False))
    project_response = client.post(
        "/api/projects",
        json={"name": "异常材料项目", "draft_text": "一种基于传感器的检测方法。"},
    )
    project_id = project_response.json()["id"]

    empty_response = client.post(
        f"/api/projects/{project_id}/materials",
        files={"file": ("empty.md", b" \n\t", "text/markdown")},
    )
    assert empty_response.status_code == 422
    assert "文件为空" in empty_response.json()["detail"]

    fake_docx_response = client.post(
        f"/api/projects/{project_id}/materials",
        files={
            "file": (
                "fake.docx",
                b"not a real docx archive",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert fake_docx_response.status_code == 422
    assert "DOCX" in fake_docx_response.json()["detail"]

    materials_response = client.get(f"/api/projects/{project_id}/materials")
    assert materials_response.status_code == 200
    assert materials_response.json()["materials"] == []


def test_disclosure_generation_ignores_failed_project_materials(tmp_path: Path):
    llm = _disclosure_llm()
    client = TestClient(
        create_app(
            data_dir=tmp_path,
            llm_client=llm,
            prior_art_provider=StaticPriorArtProvider(hits=[_prior_art_hit()]),
            load_env_file=False,
        )
    )
    project_response = client.post(
        "/api/projects",
        json={"name": "失败材料隔离", "draft_text": "一种基于图像采集的检测方法。"},
    )
    project_id = project_response.json()["id"]
    client.app.state.store.add_project_material(
        ProjectMaterial(
            id="failed-material",
            project_id=project_id,
            file_name="unsupported-round5.xyz",
            path=str(tmp_path / "unsupported-round5.xyz"),
            file_type="xyz",
            text="",
            status="failed",
            warnings=["Unsupported project material file type: .xyz"],
        )
    )
    client.app.state.store.add_project_material(
        ProjectMaterial(
            id="processed-material",
            project_id=project_id,
            file_name="valid.md",
            path=str(tmp_path / "valid.md"),
            file_type="md",
            text="有效材料：采集模块和缺陷定位模块协同工作。",
            status="processed",
            warnings=[],
        )
    )

    run_response = client.post(f"/api/projects/{project_id}/disclosures", json={"trace": False})

    assert run_response.status_code == 200
    prompts = {call.stage: call.user_prompt for call in llm.calls}
    assert "valid.md" in prompts["disclosure_scan"]
    assert "unsupported-round5.xyz" not in prompts["disclosure_scan"]
    assert "Unsupported project material file type" not in prompts["disclosure_scan"]


def test_repeated_disclosure_docx_export_does_not_mutate_export_warnings(tmp_path: Path, monkeypatch):
    monkeypatch.setattr("backend.app.disclosure.exporter.shutil.which", lambda _command: None)
    monkeypatch.delenv("PATENTS_AGENT_ENABLE_NPX_MERMAID", raising=False)
    candidate = _user_candidate("user-point", "遮挡洞口语义补全", selected=True).model_copy(
        update={
            "claim_chart": [
                ClaimChartItem(
                    prior_art_id="h1",
                    prior_art_title="一种图像缺陷检测方法",
                    differentiating_features=["多视角互证补全遮挡边界"],
                    claim_drafting_advice="写入从属权利要求。",
                )
            ]
        }
    )
    package = DisclosurePackage(
        title="图像缺陷识别交底书",
        summary="前置材料摘要",
        materials_summary="材料覆盖采集模块",
        candidates=[candidate],
        selected_candidate_id="user-point",
        prior_art_hits=[_prior_art_hit()],
        prior_art_differences="区别在遮挡洞口补全。",
        body_markdown="技术方案正文",
        mermaid="flowchart TD\nA[采集] --> B[输出]",
        image_prompt="黑白线稿。",
        export_warnings=["preexisting warning"],
    )

    export_disclosure_docx(package, tmp_path / "first.docx", tmp_path / "run-1")
    second_path = export_disclosure_docx(package, tmp_path / "second.docx", tmp_path / "run-2")

    assert package.export_warnings == ["preexisting warning"]
    doc = Document(second_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert text.count("Mermaid renderer unavailable or failed; DOCX keeps Mermaid code as text.") == 1
    assert "护城河与证据状态" in text
    assert "Claim Chart" in text
    assert "遮挡洞口语义补全" in text
    assert "写入从属权利要求。" in text


def test_disclosure_generation_fails_closed_without_llm(tmp_path: Path, monkeypatch):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    client = TestClient(
        create_app(
            data_dir=tmp_path,
            prior_art_provider=StaticPriorArtProvider(),
            load_env_file=False,
        )
    )
    project_response = client.post(
        "/api/projects",
        json={"name": "未配置模型", "draft_text": "一种AI方法，用于生成交底书。"},
    )
    project_id = project_response.json()["id"]
    material_response = client.post(
        f"/api/projects/{project_id}/materials",
        files={"file": ("material.txt", b"local material", "text/plain")},
    )
    assert material_response.status_code == 200

    disclosure_response = client.post(f"/api/projects/{project_id}/disclosures", json={"trace": False})

    assert disclosure_response.status_code == 503
    assert "DEEPSEEK_API_KEY" in disclosure_response.json()["detail"]


def _prior_art_hit() -> PriorArtHit:
    return PriorArtHit(
        id="h1",
        source="Google Patents",
        query="图像 缺陷",
        title="一种图像缺陷检测方法",
        publication_number="CN123456789A",
        url="https://patents.google.com/patent/CN123456789A",
        abstract="公开了一种图像缺陷检测方法。",
    )


def _user_candidate(candidate_id: str, title: str, selected: bool = False) -> PatentPointCandidate:
    return PatentPointCandidate(
        id=candidate_id,
        title=title,
        technical_problem="遮挡导致洞口漏识别",
        innovation="多视角互证补全洞口",
        technical_solution="融合点云和图像补全洞口边界。",
        evidence_status="feasible_unverified",
        source_type="user",
        selected=selected,
    )


def _disclosure_llm(extra: dict[str, str] | None = None) -> FakeLLMClient:
    responses = {
        "disclosure_scan": '{"summary":"图像缺陷识别项目","materials_summary":"材料覆盖采集和定位模块","technical_keywords":["图像","缺陷","神经网络"],"implementation_gaps":[]}',
        "patent_points": '{"candidates":[{"id":"p1","title":"图像缺陷识别方法及系统","technical_problem":"人工检测效率低","innovation":"基于神经网络输出缺陷位置","technical_solution":"采集图像、训练模型并输出缺陷位置","beneficial_effects":["提高检测效率"],"protection_focus":["方法","系统"],"grantability_score":0.82,"rationale":"技术链条完整"}],"selected_candidate_id":"p1"}',
        "prior_art_terms": '["图像 缺陷 神经网络","缺陷 定位 输出"]',
        "prior_art_relevance": '{"prior_art_differences":"区别在于输出缺陷位置并形成闭环审核。","hits":[{"id":"h1","relevance_summary":"涉及图像缺陷检测。","differentiators":["缺少闭环审核"]}]}',
        "disclosure_body": "# 注意事项\n本交底书用于代理人审查。\n# 一、相关技术背景\n1.1 现有技术见公开URL。\n# 三、详细技术方案\n采集图像并输出缺陷位置。",
        "disclosure_mermaid": "flowchart TD\nA[采集图像] --> B[输出缺陷位置]",
        "disclosure_image_prompt": "黑白线稿，展示图像采集到缺陷输出。",
        "disclosure_self_check": '[{"category":"URL","severity":"low","message":"现有技术URL已列出。","suggestion":"提交前人工复核链接。"}]',
    }
    responses.update(extra or {})
    return FakeLLMClient(responses)


def _create_completed_deliberation(client: TestClient, project_id: str) -> None:
    stages = _strict_deliberation_stages()
    client.app.state.store.create_deliberation_run(
        DeliberationRun(
            id=f"delib-{project_id}",
            project_id=project_id,
            status="completed",
            providers=["codex", "deepseek", "claude"],
            run_mode="full",
            stage_results=stages,
            strategy_brief=PatentStrategyBrief(
                summary="测试会审策略",
                claim_strategy=["方法独权"],
                description_strategy=["补充实施例"],
                risk_controls=["避免功能性概括"],
                agent_consensus="测试会审通过。",
            ),
            events=["test deliberation completed"],
        )
    )


def _strict_deliberation_stages() -> list[DeliberationStageResult]:
    return [
        *[
            DeliberationStageResult(
                phase="opening",
                provider_id=provider,
                label=f"opening {provider}",
                payload={"stance": "ok"},
                status="completed",
            )
            for provider in ["codex", "deepseek", "claude"]
        ],
        *[
            DeliberationStageResult(
                phase="pair",
                provider_id="codex",
                label=label,
                payload={"resolved_recommendation": "ok"},
                status="completed",
            )
            for label in ["pair codex-vs-deepseek", "pair codex-vs-claude", "pair deepseek-vs-claude"]
        ],
        DeliberationStageResult(
            phase="chair",
            provider_id="codex",
            label="chair synthesis",
            payload={"summary": "ok"},
            status="completed",
        ),
    ]
