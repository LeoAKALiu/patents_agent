from __future__ import annotations

import hashlib
import os
from pathlib import Path

from fastapi.testclient import TestClient
from docx import Document

from backend.app.external_drafts import create_external_draft_source, parse_external_draft_source
from backend.app.llm import FakeLLMClient
from backend.app.main import create_app
from backend.app.official_compile import source_draft_hash
from backend.app.schemas import DraftPackage, ProjectRecord
from backend.app.storage import SQLiteStore


def test_store_persists_external_draft_sources_and_intake_runs(tmp_path):
    db_path = tmp_path / "patents_agent.sqlite3"
    store = SQLiteStore(db_path)
    source = create_external_draft_source(
        project_id="project-1",
        source_type="pasted_text",
        text="权利要求书\n1. 一种方法。\n说明书\n本发明涉及数据处理。",
        file_name="pasted.txt",
    )
    stored_source = store.create_external_draft_source(source)
    run = parse_external_draft_source(project_id="project-1", source=stored_source)
    stored_run = store.create_external_draft_intake_run(run)

    loaded_source = store.get_external_draft_source("project-1", source.id)
    assert loaded_source is not None
    assert loaded_source.content_hash == source.content_hash
    assert store.list_external_draft_sources("project-1")[0].id == source.id

    loaded_run = store.get_external_draft_intake_run("project-1", run.id)
    assert loaded_run is not None
    assert loaded_run.working_draft_hash == run.working_draft_hash
    assert store.list_external_draft_intake_runs("project-1", source.id)[0].id == stored_run.id
    assert store.list_external_draft_intake_runs("project-1")[0].id == stored_run.id

    updated = stored_run.model_copy(update={"status": "completed"})
    returned = store.update_external_draft_intake_run(updated)
    assert returned is not None
    assert returned.status == "completed"
    assert store.get_external_draft_intake_run("project-1", run.id).status == "completed"
    missing = stored_run.model_copy(update={"id": "missing-run"})
    assert store.update_external_draft_intake_run(missing) is None

    reopened = SQLiteStore(db_path)
    reopened_source = reopened.get_external_draft_source("project-1", source.id)
    reopened_run = reopened.get_external_draft_intake_run("project-1", run.id)
    assert reopened_source is not None
    assert reopened_run is not None
    assert reopened_source.raw_text == source.raw_text
    assert reopened_run.status == "completed"


def test_delete_project_removes_external_draft_intake_records(tmp_path):
    store = SQLiteStore(tmp_path / "patents_agent.sqlite3")
    project = store.create_project(
        ProjectRecord(id="project-1", name="外部稿项目", draft_text="一种数据处理方法。")
    )
    source = store.create_external_draft_source(
        create_external_draft_source(
            project_id=project.id,
            source_type="pasted_text",
            text="权利要求书\n1. 一种方法。\n说明书\n本发明涉及数据处理。",
            file_name="pasted.txt",
        )
    )
    run = store.create_external_draft_intake_run(parse_external_draft_source(project_id=project.id, source=source))

    assert store.delete_project(project.id) is True
    assert store.get_external_draft_source(project.id, source.id) is None
    assert store.get_external_draft_intake_run(project.id, run.id) is None


def test_external_draft_api_creates_source_runs_intake_and_confirms_package(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "外部稿项目", "draft_text": "外部初稿导入项目。"},
    ).json()
    client.app.state.store.update_project_package(
        project["id"],
        DraftPackage(
            title="原始工作稿",
            abstract="原始摘要",
            claims="1. 一种原始方法，其特征在于，包括旧步骤。",
            description="原始说明书。",
            drawing_description="图1。",
            mermaid="flowchart TD\nA-->B",
            image_prompt="黑白线稿",
        ),
    )

    source_response = client.post(
        f"/api/projects/{project['id']}/external-drafts",
        json={
            "source_type": "pasted_text",
            "file_name": "draft.txt",
            "text": (
                "发明名称\n一种外部稿处理方法\n"
                "说明书\n本发明涉及专利初稿处理。\n"
            ),
        },
    )
    assert source_response.status_code == 200
    source = source_response.json()
    assert source["content_hash"]

    intake_response = client.post(f"/api/projects/{project['id']}/external-drafts/{source['id']}/intake-runs")
    assert intake_response.status_code == 200
    intake = intake_response.json()
    assert intake["status"] == "needs_review"
    assert intake["parsed_package"]["title"] == "一种外部稿处理方法"

    blank_confirm_response = client.post(
        f"/api/projects/{project['id']}/external-draft-intake-runs/{intake['id']}/confirm",
        json={
            "title": "一种外部稿处理方法",
            "abstract": "本发明公开一种外部稿处理方法。",
            "claims": "   ",
            "description": "本发明涉及专利初稿处理。",
            "drawing_description": "图1为外部稿处理流程图。",
        },
    )
    assert blank_confirm_response.status_code == 422

    confirm_response = client.post(
        f"/api/projects/{project['id']}/external-draft-intake-runs/{intake['id']}/confirm",
        json={
            "title": "一种外部稿处理方法",
            "abstract": "本发明公开一种外部稿处理方法。",
            "claims": "1. 一种外部稿处理方法，其特征在于，解析外部专利初稿并生成工作稿。",
            "description": "本发明涉及专利初稿处理。系统保存原始稿并生成内部工作稿。",
            "drawing_description": "图1为外部稿处理流程图。",
        },
    )
    assert confirm_response.status_code == 200
    confirmed = confirm_response.json()
    assert confirmed["status"] == "completed"
    assert confirmed["working_draft_hash"]

    project_after = client.get(f"/api/projects/{project['id']}").json()
    assert project_after["package"]["title"] == "一种外部稿处理方法"
    assert "保存原始稿" in project_after["package"]["description"]

    list_sources = client.get(f"/api/projects/{project['id']}/external-drafts").json()
    assert list_sources["sources"][0]["id"] == source["id"]

    list_runs = client.get(f"/api/projects/{project['id']}/external-drafts/{source['id']}/intake-runs").json()
    assert list_runs["runs"][0]["id"] == intake["id"]

    get_run_response = client.get(f"/api/projects/{project['id']}/external-draft-intake-runs/{intake['id']}")
    assert get_run_response.status_code == 200
    assert get_run_response.json()["id"] == intake["id"]

    ledger_response = client.get(f"/api/projects/{project['id']}/revision-ledger")
    assert ledger_response.status_code == 200
    records = ledger_response.json()
    assert len(records) == 1
    assert records[0]["revision_kind"] == "material_merge"
    assert records[0]["protection_scope_changed"] is True
    assert records[0]["artifact_refs"] == [f"external-draft-intake:{intake['id']}"]


def test_external_draft_docx_upload_creates_source(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "DOCX外部稿", "draft_text": "DOCX外部稿导入。"},
    ).json()
    docx_path = tmp_path / "external.docx"
    document = Document()
    document.add_paragraph("权利要求书")
    document.add_paragraph("1. 一种DOCX导入方法。")
    document.add_paragraph("说明书")
    document.add_paragraph("本发明涉及DOCX解析。")
    document.save(docx_path)

    response = client.post(
        f"/api/projects/{project['id']}/external-drafts/upload",
        files={
            "file": (
                "external.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["source_type"] == "docx_file"
    assert "DOCX导入方法" in response.json()["raw_text"]


def test_external_draft_docx_upload_rejects_malformed_docx(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "损坏DOCX外部稿", "draft_text": "损坏DOCX外部稿导入。"},
    ).json()

    response = client.post(
        f"/api/projects/{project['id']}/external-drafts/upload",
        files={
            "file": (
                "broken.docx",
                b"not a real docx package",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 422
    list_response = client.get(f"/api/projects/{project['id']}/external-drafts")
    assert list_response.status_code == 200
    assert list_response.json()["sources"] == []


def test_external_draft_upload_rejects_unsupported_file_type(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "PDF外部稿", "draft_text": "PDF外部稿导入。"},
    ).json()

    response = client.post(
        f"/api/projects/{project['id']}/external-drafts/upload",
        files={"file": ("external.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )

    assert response.status_code == 415
    list_response = client.get(f"/api/projects/{project['id']}/external-drafts")
    assert list_response.status_code == 200
    assert list_response.json()["sources"] == []


def test_external_draft_confirmed_package_runs_quality_and_bundle_report(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "外部稿提质", "draft_text": "外部初稿导入项目。"},
    ).json()
    client.app.state.store.update_project_package(
        project["id"],
        DraftPackage(
            title="一种外部稿提质方法",
            abstract="本发明公开一种外部稿提质方法。",
            claims="1. 一种方法，其特征在于，接收外部初稿并生成修改建议。",
            description="本实施例接收外部初稿。",
            drawing_description="图1为流程图。",
            mermaid="",
            image_prompt="",
            review_findings=[],
            citations=[],
            generation_logs=["external_draft_intake: confirmed from run run-1"],
        ),
    )

    completion_response = client.post(f"/api/projects/{project['id']}/completion-runs")
    assert completion_response.status_code == 200

    score_response = client.post(f"/api/projects/{project['id']}/score-improvement", json={"max_rounds": 1})
    assert score_response.status_code == 200

    compile_response = client.post(f"/api/projects/{project['id']}/official-compile-runs", json={})
    assert compile_response.status_code == 200
    assert compile_response.json()["status"] in {"completed", "blocked"}

    report_response = client.get(f"/api/projects/{project['id']}/external-draft-review-bundle/report.md")
    assert report_response.status_code == 200
    assert "EXTERNAL_DRAFT_REVIEW_BUNDLE" in report_response.text
    assert "initial_score" in report_response.text
    assert "official_compile_run_id" in report_response.text


def test_confirmed_external_draft_still_requires_official_export_gate(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, llm_client=_passing_post_draft_review_llm(), load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "外部稿正式门禁", "draft_text": "外部稿导入后仍需正式门禁。"},
    ).json()
    source = client.post(
        f"/api/projects/{project['id']}/external-drafts",
        json={
            "source_type": "pasted_text",
            "file_name": "needs-confirm.txt",
            "text": (
                "发明名称\n一种外部稿正式门禁验证方法\n"
                "说明书\n本发明涉及专利初稿处理，保存原始外部稿并生成内部工作稿。\n"
            ),
        },
    ).json()
    intake = client.post(f"/api/projects/{project['id']}/external-drafts/{source['id']}/intake-runs").json()
    assert intake["status"] == "needs_review"

    confirm_response = client.post(
        f"/api/projects/{project['id']}/external-draft-intake-runs/{intake['id']}/confirm",
        json={
            "title": "一种外部稿正式门禁验证方法",
            "abstract": "本发明公开一种外部稿正式门禁验证方法。",
            "claims": "1. 一种外部稿正式门禁验证方法，其特征在于，保存外部稿并生成内部工作稿。",
            "description": "本发明涉及专利初稿处理，保存原始外部稿并生成内部工作稿。",
            "drawing_description": "图1为外部稿正式门禁验证流程图。",
        },
    )
    assert confirm_response.status_code == 200

    blocked_without_compile = client.get(f"/api/projects/{project['id']}/official-export.md")
    assert blocked_without_compile.status_code == 409
    assert "Quality checks are required" in blocked_without_compile.json()["detail"]
    assert client.post(f"/api/projects/{project['id']}/filing-readiness").status_code == 200
    assert client.post(f"/api/projects/{project['id']}/claim-defense-worksheets").status_code == 200
    assert client.post(f"/api/projects/{project['id']}/completion-runs").status_code == 200

    blocked_without_compile = client.get(f"/api/projects/{project['id']}/official-export.md")
    assert blocked_without_compile.status_code == 409
    assert "Official draft compile is required" in blocked_without_compile.json()["detail"]

    compile_response = client.post(f"/api/projects/{project['id']}/official-compile-runs", json={})
    assert compile_response.status_code == 200
    assert compile_response.json()["status"] == "completed"

    blocked_without_review = client.get(f"/api/projects/{project['id']}/official-export.md")
    assert blocked_without_review.status_code == 409
    assert "Post-draft multi-agent review is required" in blocked_without_review.json()["detail"]

    review_response = client.post(f"/api/projects/{project['id']}/post-draft-reviews", json={})
    assert review_response.status_code == 200
    assert review_response.json()["export_allowed"] is True

    export_response = client.get(f"/api/projects/{project['id']}/official-export.md")
    assert export_response.status_code == 200
    assert "一种外部稿正式门禁验证方法" in export_response.text
    assert "权利要求书" in export_response.text
    assert "external_draft_intake" not in export_response.text
    assert "EXTERNAL_DRAFT_REVIEW_BUNDLE" not in export_response.text
    assert "attorney_memo" not in export_response.text


def test_external_draft_review_bundle_pairs_intake_run_with_its_source(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "外部稿配对", "draft_text": "外部初稿导入项目。"},
    ).json()

    source_with_run = client.post(
        f"/api/projects/{project['id']}/external-drafts",
        json={
            "source_type": "pasted_text",
            "file_name": "with-run.txt",
            "text": (
                "发明名称\n一种外部稿处理方法\n"
                "权利要求书\n1. 一种外部稿处理方法。\n"
                "说明书\n本发明涉及专利初稿处理。\n"
            ),
        },
    ).json()
    intake = client.post(
        f"/api/projects/{project['id']}/external-drafts/{source_with_run['id']}/intake-runs"
    ).json()
    newer_source = client.post(
        f"/api/projects/{project['id']}/external-drafts",
        json={
            "source_type": "pasted_text",
            "file_name": "newer-without-run.txt",
            "text": "发明名称\n一种后续外部稿\n说明书\n本发明涉及后续稿。",
        },
    ).json()

    report_response = client.get(f"/api/projects/{project['id']}/external-draft-review-bundle/report.md")

    assert report_response.status_code == 200
    assert f"- source_id: {source_with_run['id']}" in report_response.text
    assert f"- intake_run_id: {intake['id']}" in report_response.text
    assert f"- source_id: {newer_source['id']}" not in report_response.text


def _passing_post_draft_review_llm() -> FakeLLMClient:
    return FakeLLMClient(
        {
            "post_draft_claims_reviewer": """
{
  "role": "claims_reviewer",
  "status": "passed",
  "blocking_issues": [],
  "contamination_hits": [],
  "rewrite_suggestions": [],
  "official_safe_patches": [],
  "attorney_memo": ["代理人可继续复核权利要求。"]
}
""",
            "post_draft_spec_cleaner": """
{
  "role": "spec_cleaner",
  "status": "passed",
  "blocking_issues": [],
  "contamination_hits": [],
  "rewrite_suggestions": [],
  "official_safe_patches": [],
  "attorney_memo": ["清污检查通过。"]
}
""",
            "post_draft_technical_hardness": """
{
  "role": "technical_hardness",
  "status": "passed",
  "blocking_issues": [],
  "contamination_hits": [],
  "rewrite_suggestions": [],
  "official_safe_patches": [],
  "attorney_memo": ["技术硬度可继续增强。"]
}
""",
            "post_draft_chair_synthesis": """
{
  "status": "passed",
  "export_allowed": true,
  "blocking_issues": [],
  "contamination_hits": [],
  "claim_1_rewrite": "",
  "system_claim_rewrite": "",
  "abstract_rewrite": "",
  "description_rewrite_tasks": [],
  "official_safe_patches": [],
  "attorney_memo": ["主席综合意见仅保留在内部报告。"],
  "next_actions": []
}
""",
        }
    )


def test_docx_upload_stores_original_sealed_and_hashes_match(tmp_path):
    """DOCX import writes the original to disk read-only and records a hash
    that can be cross-checked against the file on disk for tampering."""
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "DOCX密封验证", "draft_text": "DOCX导入并密封原始文件。"},
    ).json()

    docx_path = tmp_path / "sealed.docx"
    document = Document()
    document.add_paragraph("权利要求书")
    document.add_paragraph("1. 一种DOCX密封方法，其特征在于，导入外部稿并锁定原始文件。")
    document.add_paragraph("说明书")
    document.add_paragraph("本发明涉及DOCX外部稿的密封保存。")
    document.save(docx_path)
    original_bytes = docx_path.read_bytes()

    response = client.post(
        f"/api/projects/{project['id']}/external-drafts/upload",
        files={
            "file": (
                "sealed.docx",
                original_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["source_type"] == "docx_file"
    assert payload["raw_path"], "raw_path should be persisted for sealed DOCX imports"
    assert payload["metadata"]["raw_content_hash"], "raw_content_hash should be recorded"

    sealed_path = Path(payload["raw_path"])
    assert sealed_path.exists()
    sealed_bytes = sealed_path.read_bytes()
    assert sealed_bytes == original_bytes

    # The raw bytes hash recorded in the source must match the on-disk file
    # verbatim — this is the integrity check that backs the read-only seal.
    assert payload["metadata"]["raw_content_hash"] == hashlib.sha256(sealed_bytes).hexdigest()

    # The seal is best-effort: on POSIX, the file should drop its write bits.
    if os.name != "nt":
        mode = sealed_path.stat().st_mode & 0o777
        assert mode & 0o222 == 0, f"Expected read-only mode, got {oct(mode)}"

    # The normalized text hash (the one tied to the parsed working draft) is
    # derived from the docx text content, not the raw bytes.
    assert payload["content_hash"] != payload["metadata"]["raw_content_hash"]


def test_official_compile_is_idempotent_on_imported_external_drafts(tmp_path):
    """Running the official compile twice on the same imported draft must
    produce identical contamination_removed, blocked_items, and the same
    official_package_hash — guaranteeing the contamination check is stable
    and doesn't accumulate cruft across re-runs."""
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "外部稿正式稿幂等", "draft_text": "外部稿导入。"},
    ).json()

    # Import a draft via the API and confirm it through the normal pipeline.
    source = client.post(
        f"/api/projects/{project['id']}/external-drafts",
        json={
            "source_type": "pasted_text",
            "file_name": "idempotent.txt",
            "text": (
                "发明名称\n一种外部稿正式稿幂等测试方法\n"
                "权利要求书\n1. 一种外部稿正式稿幂等测试方法，其特征在于，"
                "重新运行正式编译不引入新污染。\n"
                "说明书\n本发明涉及对外部稿的正式稿编译做幂等性检查。"
            ),
        },
    ).json()
    intake = client.post(
        f"/api/projects/{project['id']}/external-drafts/{source['id']}/intake-runs"
    ).json()
    client.post(
        f"/api/projects/{project['id']}/external-draft-intake-runs/{intake['id']}/confirm",
        json={
            "title": "一种外部稿正式稿幂等测试方法",
            "abstract": "本发明公开一种外部稿正式稿幂等测试方法。",
            "claims": (
                "1. 一种外部稿正式稿幂等测试方法，其特征在于，"
                "重新运行正式编译不引入新污染。"
            ),
            "description": "本发明涉及对外部稿的正式稿编译做幂等性检查。",
            "drawing_description": "图1为正式稿编译幂等测试流程图。",
        },
    )

    # First run establishes the baseline.
    first = client.post(
        f"/api/projects/{project['id']}/official-compile-runs", json={}
    ).json()
    assert first["status"] == "completed"

    # Second run must agree on the contamination and hash story.
    second = client.post(
        f"/api/projects/{project['id']}/official-compile-runs", json={}
    ).json()
    assert second["status"] == "completed"

    assert first["source_draft_hash"] == second["source_draft_hash"]
    assert first["official_package_hash"] == second["official_package_hash"]
    assert first["contamination_removed"] == second["contamination_removed"]
    assert first["blocked_items"] == second["blocked_items"]

    # The working draft on the project side must match the source_draft_hash
    # the official compile computed, proving the import + compile pair is closed.
    project_state = client.get(f"/api/projects/{project['id']}").json()
    draft_hash = source_draft_hash(DraftPackage.model_validate(project_state["package"]))
    assert draft_hash == first["source_draft_hash"]


def test_official_compile_blocks_contaminated_imported_draft(tmp_path):
    """If a confirmation smuggles internal drafting text into the imported
    draft, the official compile must still flag it via the contamination
    check — proving the import path does not weaken the export gate."""
    client = TestClient(create_app(data_dir=tmp_path, load_env_file=False))
    project = client.post(
        "/api/projects",
        json={"name": "外部稿污染门禁", "draft_text": "外部稿污染门禁验证。"},
    ).json()

    source = client.post(
        f"/api/projects/{project['id']}/external-drafts",
        json={
            "source_type": "pasted_text",
            "file_name": "smuggled.txt",
            "text": (
                "发明名称\n一种外部稿污染门禁验证方法\n"
                "权利要求书\n1. 一种方法。\n说明书\n本发明涉及外部稿污染门禁。"
            ),
        },
    ).json()
    intake = client.post(
        f"/api/projects/{project['id']}/external-drafts/{source['id']}/intake-runs"
    ).json()
    # Confirm with claims that have been seeded with internal drafting noise
    # the same way the generator sometimes leaks it.
    client.post(
        f"/api/projects/{project['id']}/external-draft-intake-runs/{intake['id']}/confirm",
        json={
            "title": "一种外部稿污染门禁验证方法",
            "abstract": "本发明公开一种外部稿污染门禁验证方法。",
            "claims": "好的，下面撰写权利要求书。\n1. 一种外部稿污染门禁验证方法。",
            "description": "本发明涉及外部稿污染门禁。support_gap: 仍需补实验。",
            "drawing_description": "图1为流程图。",
        },
    )

    blocked = client.post(
        f"/api/projects/{project['id']}/official-compile-runs", json={}
    ).json()
    assert blocked["status"] == "blocked"
    assert any(
        item.get("category") in {"residual_internal_text", "empty_required_section"}
        for item in blocked["blocked_items"]
    )
    assert blocked["official_package"] is None
    # The official export gate must not let the contaminated imported draft
    # through, even though it cleared the intake.
    blocked_export = client.get(f"/api/projects/{project['id']}/official-export.md")
    assert blocked_export.status_code == 409
