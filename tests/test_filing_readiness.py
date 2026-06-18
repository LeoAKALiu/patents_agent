from docx import Document
from fastapi.testclient import TestClient

from backend.app.schemas import FilingReadinessIssue, FilingReadinessReport
from backend.app.filing_readiness import (
    assess_filing_readiness,
    export_official_docx,
    official_package_to_markdown,
    readiness_report_to_markdown,
)
from backend.app.llm import FakeLLMClient
from backend.app.main import create_app
from backend.app.schemas import DraftPackage, PatentStrategyBrief


def test_filing_readiness_report_status_and_issue_shape():
    issue = FilingReadinessIssue(
        category="internal_trace",
        severity="high",
        target="claims",
        matched_text="根据会审策略撰写",
        message="正式稿包含内部过程痕迹。",
        suggestion="删除过程性表述，仅保留权利要求内容。",
        can_auto_clean=True,
    )
    report = FilingReadinessReport(
        id="r1",
        project_id="p1",
        draft_package_hash="hash-1",
        status="high_risk",
        rules_version="filing-readiness-v1",
        issues=[issue],
    )

    assert report.status == "high_risk"
    assert report.issues[0].category == "internal_trace"
    assert report.issues[0].target == "claims"


def _dirty_package() -> DraftPackage:
    return DraftPackage(
        title="一种外立面逆建模方法",
        abstract="本发明公开了一种方法，效率提升30%。",
        claims="根据会审策略撰写\n```mermaid\nflowchart TD\nA-->B\n```",
        description="本发明属于人工智能软件方法领域，可能不具备创造性。",
        drawing_description="图1为流程图。",
        mermaid="flowchart TD\nA[点云] --> B[IFC]",
        image_prompt="黑白线稿 prompt",
        review_findings=[],
        citations=[],
        generation_logs=["generation_logs: claims, description, image_prompt"],
        strategy_brief=PatentStrategyBrief(
            summary="多Agent会审指出存在充分公开风险。",
            claim_strategy=["根据技术交底书补强独权。"],
            description_strategy=[],
            risk_controls=[],
            agent_consensus="deliberation complete",
        ),
    )


def _clean_base_package(**overrides) -> DraftPackage:
    data = {
        "title": "一种外立面逆建模方法",
        "abstract": "本发明公开了一种外立面逆建模方法。",
        "claims": "1. 一种外立面逆建模方法，其特征在于包括点云处理步骤。",
        "description": "本发明涉及建筑外立面数据处理技术。",
        "drawing_description": "图1为方法流程图。",
        "mermaid": "",
        "image_prompt": "",
        "review_findings": [],
        "citations": [],
        "generation_logs": [],
    }
    data.update(overrides)
    return DraftPackage(**data)


def test_clean_gate_detects_internal_and_format_pollution():
    report = assess_filing_readiness("project-1", _dirty_package(), verified_effects=False)

    categories = {issue.category for issue in report.issues}
    assert report.status == "high_risk"
    assert "format_pollution" in categories
    assert "internal_trace" in categories
    assert "unfavorable_statement" in categories
    assert "unverified_effect" in categories
    assert "subject_matter_risk" in categories


def test_readiness_report_markdown_contains_matches_and_suggestions():
    report = assess_filing_readiness("project-1", _dirty_package(), verified_effects=False)
    markdown = readiness_report_to_markdown(report)

    assert "# FILING_READINESS_REPORT" in markdown
    assert "根据会审策略撰写" in markdown
    assert "正式稿包含内部过程痕迹" in markdown
    assert "建筑信息模型、三维点云处理" in markdown


def test_official_markdown_cleans_internal_and_format_pollution_lines():
    markdown = official_package_to_markdown(_dirty_package())

    assert "```" not in markdown
    assert "flowchart TD" not in markdown
    assert "根据会审策略" not in markdown
    assert "prompt" not in markdown
    assert "generation_logs" not in markdown
    assert "## 摘要" in markdown
    assert "## 权利要求书" in markdown
    assert "## 说明书" in markdown
    assert "## 附图说明" in markdown


def test_only_unverified_quantitative_effect_warns_with_medium_severity():
    package = _clean_base_package(abstract="本发明公开了一种方法，效率提升30%。")

    report = assess_filing_readiness("project-1", package, verified_effects=False)

    assert report.status == "warning"
    assert {issue.category for issue in report.issues} == {"unverified_effect"}
    assert {issue.severity for issue in report.issues} == {"medium"}


def test_filing_readiness_allows_technical_disclosure_reference_phrase():
    package = _clean_base_package(
        description=(
            "根据技术交底书中记载的传感器布置方式，控制器获取第一检测信号和第二检测信号，"
            "并基于二者的匹配关系生成告警指令。"
        )
    )

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert not any(
        issue.category == "internal_trace" and "根据技术交底书" in issue.matched_text
        for issue in report.issues
    )


def test_filing_readiness_ignores_prior_art_percent_effect_context():
    package = _clean_base_package(
        description=(
            "对比文件CN123456公开的方案可提升15%识别速度，但该方案依赖云端集中推理。"
            "本实施例采用边缘端缓存策略降低通信等待时间。"
        )
    )

    report = assess_filing_readiness("project-1", package, verified_effects=False)

    assert not any(issue.category == "unverified_effect" for issue in report.issues)


def test_filing_readiness_flags_current_invention_percent_after_prior_art_sentence():
    package = _clean_base_package(
        description="对比文件CN123456公开了云端推理。本实施例可提升15%识别速度。"
    )

    report = assess_filing_readiness("project-1", package, verified_effects=False)

    assert any(
        issue.category == "unverified_effect" and issue.matched_text == "提升15%"
        for issue in report.issues
    )


def test_filing_readiness_flags_solution_percent_after_prior_art_problem_clause():
    package = _clean_base_package(
        description="现有技术记载存在云端延迟问题，提出一种边缘缓存策略，可提升15%识别效率。"
    )

    report = assess_filing_readiness("project-1", package, verified_effects=False)

    assert any(
        issue.category == "unverified_effect" and issue.matched_text == "提升15%"
        for issue in report.issues
    )


def test_filing_readiness_allows_defensive_unverified_example_wording():
    package = _clean_base_package(
        description=(
            "尚未验证的参数组合仅作为可选实施例，不用于限定本申请的技术效果。"
        )
    )

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert not any(
        issue.category == "unfavorable_statement" and "尚未验证" in issue.matched_text
        for issue in report.issues
    )


def test_filing_readiness_flags_direct_unverified_effect_admission():
    package = _clean_base_package(description="本申请的技术效果尚未验证。")

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert any(
        issue.category == "unfavorable_statement" and "尚未验证" in issue.matched_text
        for issue in report.issues
    )


def test_filing_readiness_blocks_technical_disclosure_drafting_process_phrase():
    package = _clean_base_package(description="根据技术交底书进行撰写权利要求书。")

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert any(
        issue.category == "internal_trace" and "根据技术交底书" in issue.matched_text
        for issue in report.issues
    )


def test_filing_readiness_blocks_technical_disclosure_object_before_drafting_verb():
    package = _clean_base_package(description="根据技术交底书进行权利要求书撰写。")

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert any(
        issue.category == "internal_trace" and "根据技术交底书" in issue.matched_text
        for issue in report.issues
    )


def test_filing_readiness_blocks_technical_disclosure_claim_strengthening_phrase():
    package = _clean_base_package(description="根据技术交底书补强独权。")

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert any(
        issue.category == "internal_trace" and "根据技术交底书" in issue.matched_text
        for issue in report.issues
    )


def test_filing_readiness_blocks_technical_disclosure_object_before_strengthening_verb():
    package = _clean_base_package(description="根据技术交底书对独权进行补强。")

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert any(
        issue.category == "internal_trace" and "根据技术交底书" in issue.matched_text
        for issue in report.issues
    )


def test_only_subject_matter_risk_warns_with_medium_severity():
    package = _clean_base_package(description="本发明属于人工智能软件方法领域。")

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    assert report.status == "warning"
    assert {issue.category for issue in report.issues} == {"subject_matter_risk"}
    assert {issue.severity for issue in report.issues} == {"medium"}


def test_internal_field_presence_blocks_chinese_only_prompt_and_plain_logs():
    package = _clean_base_package(
        image_prompt="黑白线稿，展示立面构件与点云配准关系。",
        generation_logs=["完成权利要求和说明书初稿生成。"],
    )

    report = assess_filing_readiness("project-1", package, verified_effects=True)

    issues_by_match = {issue.matched_text: issue for issue in report.issues}
    assert report.status == "high_risk"
    assert issues_by_match["image_prompt"].category == "format_pollution"
    assert issues_by_match["image_prompt"].severity == "high"
    assert issues_by_match["generation_logs"].category == "internal_trace"
    assert issues_by_match["generation_logs"].severity == "high"


def test_official_markdown_preserves_body_when_removing_embedded_internal_trace():
    package = _clean_base_package(
        claims="1. 一种方法，包括采集点云，根据会审策略撰写，并生成建筑构件参数。"
    )

    markdown = official_package_to_markdown(package)

    assert "1. 一种方法，包括采集点云" in markdown
    assert "并生成建筑构件参数。" in markdown
    assert "根据会审策略" not in markdown


def test_official_markdown_removes_regex_internal_trace_fragments():
    package = _clean_base_package(
        claims="1. 一种方法，包括采集点云，根据技术交底书进行撰写权利要求书，并生成建筑构件参数。"
    )

    markdown = official_package_to_markdown(package)

    assert "1. 一种方法，包括采集点云" in markdown
    assert "并生成建筑构件参数。" in markdown
    assert "根据技术交底书" not in markdown


def test_export_official_docx_cleans_pollution_and_keeps_formal_sections(tmp_path):
    output_path = tmp_path / "official.docx"

    export_official_docx(_dirty_package(), output_path)

    doc = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "摘要" in text
    assert "权利要求书" in text
    assert "说明书" in text
    assert "附图说明" in text
    assert "```" not in text
    assert "flowchart TD" not in text
    assert "prompt" not in text
    assert "generation_logs" not in text
    assert "根据会审策略" not in text


def test_official_exports_do_not_emit_internal_contamination_markers(tmp_path):
    package = _dirty_package()
    output_path = tmp_path / "official-clean.docx"
    forbidden = [
        "```",
        "flowchart TD",
        "prompt",
        "generation_logs",
        "根据会审策略",
        "多Agent会审",
    ]

    markdown = official_package_to_markdown(package)
    export_official_docx(package, output_path)
    doc = Document(output_path)
    docx_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    for marker in forbidden:
        assert marker not in markdown
        assert marker not in docx_text


from backend.app.storage import SQLiteStore


def test_store_persists_multiple_filing_readiness_reports(tmp_path):
    store = SQLiteStore(tmp_path / "patents_agent.sqlite3")
    first = FilingReadinessReport(id="r1", project_id="p1", status="warning", issues=[])
    second = FilingReadinessReport(id="r2", project_id="p1", status="high_risk", issues=[])

    store.create_filing_readiness_report(first)
    store.create_filing_readiness_report(second)

    reports = store.list_filing_readiness_reports("p1")
    assert [report.id for report in reports] == ["r2", "r1"]
    assert store.get_filing_readiness_report("p1", "r1").status == "warning"


def test_filing_readiness_api_no_longer_unlocks_official_export_without_post_draft_review(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, llm_client=FakeLLMClient({}), load_env_file=False))
    project_id = client.post(
        "/api/projects",
        json={"name": "清稿测试", "draft_text": "一种建筑外立面逆建模方法。"},
    ).json()["id"]
    dirty_package = _dirty_package()
    client.app.state.store.update_project_package(project_id, dirty_package)

    report_response = client.post(f"/api/projects/{project_id}/filing-readiness")
    assert report_response.status_code == 200
    report = report_response.json()
    assert report["status"] == "high_risk"

    list_response = client.get(f"/api/projects/{project_id}/filing-readiness")
    assert list_response.json()["reports"][0]["id"] == report["id"]

    report_md = client.get(f"/api/projects/{project_id}/filing-readiness/{report['id']}/export.md")
    assert report_md.status_code == 200
    assert "FILING_READINESS_REPORT" in report_md.text
    assert "根据会审策略撰写" in report_md.text

    official_md = client.get(f"/api/projects/{project_id}/official-export.md")
    assert official_md.status_code == 409
    assert "Official draft compile is required" in official_md.json()["detail"]


def test_filing_readiness_api_warning_allows_official_markdown_and_docx_export(tmp_path):
    client = TestClient(create_app(data_dir=tmp_path, llm_client=FakeLLMClient(_passed_post_draft_review_responses()), load_env_file=False))
    project_id = client.post(
        "/api/projects",
        json={"name": "警告导出测试", "draft_text": "一种建筑外立面逆建模方法。"},
    ).json()["id"]
    warning_package = _clean_base_package(abstract="本发明公开了一种方法，效率提升30%。")
    client.app.state.store.update_project_package(project_id, warning_package)

    report_response = client.post(f"/api/projects/{project_id}/filing-readiness")
    assert report_response.status_code == 200
    assert report_response.json()["status"] == "warning"

    compile_response = client.post(f"/api/projects/{project_id}/official-compile-runs", json={})
    assert compile_response.status_code == 200

    review_response = client.post(f"/api/projects/{project_id}/post-draft-reviews", json={})
    assert review_response.status_code == 200
    assert review_response.json()["export_allowed"] is True

    official_md = client.get(f"/api/projects/{project_id}/official-export.md")
    assert official_md.status_code == 200
    assert "权利要求书" in official_md.text

    official_docx = client.get(f"/api/projects/{project_id}/official-export.docx")
    assert official_docx.status_code == 200
    assert official_docx.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def _passed_post_draft_review_responses() -> dict[str, str]:
    def role_payload(role: str) -> str:
        return f"""
{{
  "role": "{role}",
  "status": "passed",
  "blocking_issues": [],
  "contamination_hits": [],
  "rewrite_suggestions": ["正式稿结构可提交。"],
  "official_safe_patches": [],
  "attorney_memo": ["代理人复核后提交。"]
}}
"""

    return {
        "post_draft_claims_reviewer": role_payload("claims_reviewer"),
        "post_draft_spec_cleaner": role_payload("spec_cleaner"),
        "post_draft_technical_hardness": role_payload("technical_hardness"),
        "post_draft_chair_synthesis": """
{
  "status": "passed",
  "export_allowed": true,
  "blocking_issues": [],
  "contamination_hits": [],
  "claim_1_rewrite": "",
  "system_claim_rewrite": "",
  "abstract_rewrite": "",
  "description_rewrite_tasks": [],
  "official_safe_patches": [],
  "attorney_memo": ["成稿会审通过。"],
  "next_actions": ["允许正式导出。"]
}
""",
    }
