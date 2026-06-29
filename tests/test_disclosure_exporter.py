from backend.app.disclosure.exporter import (
    clean_disclosure_to_markdown,
    disclosure_sidecar_to_markdown,
    export_disclosure_docx,
)
from docx import Document
from backend.app.schemas import DisclosurePackage, DisclosureSelfCheckFinding, PatentPointCandidate, PriorArtHit


def _package() -> DisclosurePackage:
    return DisclosurePackage(
        title="一种图像缺陷识别方法",
        summary="摘要",
        materials_summary="材料覆盖",
        candidates=[
            PatentPointCandidate(
                id="p1",
                title="闭环反馈",
                technical_problem="效率低",
                innovation="实时反馈",
                technical_solution="检测后回写采集策略",
                protection_focus=["方法"],
            )
        ],
        selected_candidate_id="p1",
        prior_art_hits=[
            PriorArtHit(
                id="h1",
                source="Google Patents",
                query="图像 缺陷",
                title="一种图像缺陷检测方法",
                publication_number="CN123456789A",
                url="https://patents.google.com/patent/CN123456789A",
                abstract="公开了缺陷检测。",
                differentiators=["缺少闭环反馈"],
            )
        ],
        prior_art_differences="现有技术缺少闭环反馈。",
        body_markdown="# 技术交底书正文\n\n## 一、背景\n正文。",
        mermaid="flowchart TD\nA-->B",
        image_prompt="黑白线稿",
        self_check_findings=[
            DisclosureSelfCheckFinding(category="url", severity="low", message="URL 存在", suggestion="无")
        ],
        generation_logs=["project_scan: summarized draft", "warning: internal"],
        research_ledger={"entries": [{"provider": "google_patents"}]},
        provider_diagnostics=[{"phase": "post_flight"}],
        research_confidence="medium",
    )


def test_clean_disclosure_excludes_internal_sections() -> None:
    markdown = clean_disclosure_to_markdown(_package())

    assert "技术交底书正文" in markdown
    assert "https://patents.google.com/patent/CN123456789A" in markdown
    assert "Claim Chart" not in markdown
    assert "检索来源台账" not in markdown
    assert "自检结果" not in markdown
    assert "生成日志" not in markdown
    assert "provider_diagnostics" not in markdown
    assert "evidence_id" not in markdown


def test_clean_disclosure_scrubs_internal_metadata_from_llm_body() -> None:
    package = _package().model_copy(
        update={
            "body_markdown": (
                "# 技术交底书正文\n\n"
                "## 一、技术领域\n"
                "本发明涉及图像缺陷识别。\n\n"
                "evidence_id: E-001\n"
                "research_ledger: internal provider trace\n"
                "- source_ledger: CN123456789A\n\n"
                "## Claim Chart\n"
                "| prior_art | overlap |\n"
                "| --- | --- |\n"
                "| CN123456789A | 内部对比 |\n\n"
                "## 二、发明内容\n"
                "系统根据检测结果实时回写采集策略。\n\n"
                "## 生成日志\n"
                "- disclosure_body: internal trace\n\n"
                "## 三、具体实施方式\n"
                "在一个实施例中，采集模块将反馈信号写入任务队列。"
            )
        }
    )

    markdown = clean_disclosure_to_markdown(package)

    assert "本发明涉及图像缺陷识别" in markdown
    assert "系统根据检测结果实时回写采集策略" in markdown
    assert "采集模块将反馈信号写入任务队列" in markdown
    assert "evidence_id" not in markdown
    assert "research_ledger" not in markdown
    assert "source_ledger" not in markdown
    assert "Claim Chart" not in markdown
    assert "生成日志" not in markdown


def test_clean_disclosure_scrubs_inline_and_chinese_internal_markers_from_llm_body() -> None:
    package = _package().model_copy(
        update={
            "body_markdown": (
                "# 技术交底书正文\n\n"
                "本段保留正常技术方案。\n"
                "本段包含 evidence_id: E-001，不能进入清洁交底书。\n"
                "修订记录：根据内部审查意见改写权利要求。\n\n"
                "## revision_ledger\n"
                "- changed claims\n\n"
                "## 二、实施方式\n"
                "实施方式正文保留。"
            )
        }
    )

    markdown = clean_disclosure_to_markdown(package)

    assert "本段保留正常技术方案" in markdown
    assert "实施方式正文保留" in markdown
    assert "evidence_id" not in markdown
    assert "修订记录" not in markdown
    assert "changed claims" not in markdown


def test_clean_disclosure_docx_scrubs_internal_metadata_from_llm_body(tmp_path) -> None:
    package = _package().model_copy(
        update={
            "body_markdown": (
                "# 技术交底书正文\n\n"
                "正文保留的技术方案。\n"
                "generation_logs: disclosure_body internal trace\n\n"
                "## 检索来源台账\n"
                "| provider | query |\n"
                "| --- | --- |\n"
                "| internal | secret |\n\n"
                "## 具体实施方式\n"
                "实施例保留。"
            )
        }
    )
    docx_path = export_disclosure_docx(package, tmp_path / "disclosure.docx", tmp_path)
    text = "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs)

    assert "正文保留的技术方案" in text
    assert "实施例保留" in text
    assert "generation_logs" not in text
    assert "检索来源台账" not in text
    assert "internal | secret" not in text


def test_clean_disclosure_scrubs_internal_metadata_table_rows_and_blocks() -> None:
    package = _package().model_copy(
        update={
            "body_markdown": (
                "# 技术交底书正文\n\n"
                "## 一、技术领域\n"
                "本发明涉及图像缺陷识别。\n\n"
                "| 字段 | 内容 |\n"
                "| --- | --- |\n"
                "| evidence_id | E-001 |\n"
                "| source_label | 内部样本 |\n"
                "| 字段 | evidence_id |\n"
                "| 技术效果 | 降低误检率 |\n\n"
                "## 二、发明内容\n"
                "系统根据检测结果实时回写采集策略。"
            )
        }
    )

    markdown = clean_disclosure_to_markdown(package)

    assert "本发明涉及图像缺陷识别" in markdown
    assert "降低误检率" in markdown
    assert "evidence_id" not in markdown
    assert "source_label" not in markdown
    assert "E-001" not in markdown
    assert "内部样本" not in markdown
    assert "| 字段 | evidence_id |" not in markdown


def test_clean_disclosure_scrubs_value_cell_metadata_table_rows_from_markdown_and_docx(tmp_path) -> None:
    package = _package().model_copy(
        update={
            "body_markdown": (
                "# 技术交底书正文\n\n"
                "| 字段 | 内容 |\n"
                "| --- | --- |\n"
                "| 技术效果 | 降低误检率 |\n"
                "| 备注 | evidence_id: E-001 |\n"
                "| 说明 | source_label: internal |\n"
                "| 说明 | 修订记录：内部改写 |\n\n"
                "后续技术方案保留。"
            )
        }
    )

    markdown = clean_disclosure_to_markdown(package)
    docx_path = export_disclosure_docx(package, tmp_path / "disclosure.docx", tmp_path)
    docx_text = "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs)

    for text in (markdown, docx_text):
        assert "降低误检率" in text
        assert "后续技术方案保留" in text
        assert "evidence_id" not in text
        assert "source_label" not in text
        assert "修订记录" not in text
        assert "E-001" not in text
        assert "internal" not in text


def test_clean_disclosure_scrubs_structured_internal_blocks_from_markdown_and_docx(tmp_path) -> None:
    package = _package().model_copy(
        update={
            "body_markdown": (
                "---\n"
                "title: 一种图像缺陷识别方法\n"
                "source_id: src-001\n"
                "internal_only: true\n"
                "---\n\n"
                "# 技术交底书正文\n\n"
                "```yaml\n"
                "evidence_id: E-001\n"
                "source_label: internal-dataset\n"
                "```\n\n"
                "## 一、技术领域\n"
                "本发明涉及图像缺陷识别。\n\n"
                "```json\n"
                "{\"source_label\": \"internal\", \"technical_effect\": \"降低误检率\"}\n"
                "```\n\n"
                "## 二、发明内容\n"
                "系统根据检测结果实时回写采集策略。"
            )
        }
    )

    markdown = clean_disclosure_to_markdown(package)
    docx_path = export_disclosure_docx(package, tmp_path / "disclosure.docx", tmp_path)
    docx_text = "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs)

    for text in (markdown, docx_text):
        assert "本发明涉及图像缺陷识别" in text
        assert "系统根据检测结果实时回写采集策略" in text
        assert "source_id" not in text
        assert "internal_only" not in text
        assert "evidence_id" not in text
        assert "source_label" not in text
        assert "internal-dataset" not in text


def test_sidecar_contains_internal_sections() -> None:
    markdown = disclosure_sidecar_to_markdown(_package())

    assert "Claim Chart" in markdown
    assert "检索来源台账" in markdown
    assert "自检结果" in markdown
    assert "生成日志" in markdown
    assert "Google Patents" in markdown


def test_clean_disclosure_appends_only_missing_public_prior_art_urls() -> None:
    package = _package().model_copy(
        update={
            "body_markdown": (
                "# 技术交底书正文\n\n"
                "正文已引用 https://patents.google.com/patent/CN123456789A 。"
            ),
            "prior_art_hits": [
                *_package().prior_art_hits,
                PriorArtHit(
                    id="h2",
                    source="Google Patents",
                    query="图像 缺陷",
                    title="另一篇现有技术",
                    publication_number="US20240123456A1",
                    url="https://patents.google.com/patent/US20240123456A1",
                    abstract="公开了另一种处理方式。",
                ),
                PriorArtHit(
                    id="h3",
                    source="Google Patents",
                    query="图像 缺陷",
                    title="无公开链接条目",
                    publication_number="CN000000001A",
                    url="",
                ),
            ],
        }
    )

    markdown = clean_disclosure_to_markdown(package)

    assert markdown.count("https://patents.google.com/patent/CN123456789A") == 1
    assert "https://patents.google.com/patent/US20240123456A1" in markdown
    assert "无公开链接条目" not in markdown


def test_clean_disclosure_does_not_append_unsupported_prior_art_urls() -> None:
    package = _package().model_copy(
        update={
            "body_markdown": "# 技术交底书正文\n\n正文。",
            "prior_art_hits": [
                PriorArtHit(
                    id="h1",
                    source="DeepResearch Markdown",
                    query="图像 缺陷",
                    title="非公开视频链接",
                    publication_number="CN123456789A",
                    url="https://example.com/patent/CN123456789A",
                    abstract="公开了缺陷检测。",
                ),
                PriorArtHit(
                    id="h2",
                    source="Google Patents",
                    query="图像 缺陷",
                    title="公开链接条目",
                    publication_number="US20240123456A1",
                    url="https://patents.google.com/patent/US20240123456A1",
                    abstract="公开了另一种处理方式。",
                ),
            ],
        }
    )

    markdown = clean_disclosure_to_markdown(package)

    assert "https://example.com/patent/CN123456789A" not in markdown
    assert "非公开视频链接" not in markdown
    assert "https://patents.google.com/patent/US20240123456A1" in markdown


def test_clean_disclosure_normalizes_trailing_punctuation_when_comparing_urls() -> None:
    package = _package().model_copy(
        update={
            "body_markdown": "# 技术交底书正文\n\n正文已引用 https://patents.google.com/patent/CN123456789A 。",
        }
    )

    markdown = clean_disclosure_to_markdown(package)

    assert markdown.count("https://patents.google.com/patent/CN123456789A") == 1
